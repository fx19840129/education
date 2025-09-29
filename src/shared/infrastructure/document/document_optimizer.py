#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档生成优化器
优化Word文档生成性能，支持模板化、并行处理和流式生成
"""

import os
import time
import threading
from typing import Dict, Any, List, Optional, Union, Callable
from dataclasses import dataclass, field
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import json

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..monitoring.metrics import get_metrics
from ..cache.cache_manager import get_cache
from ..errors.exceptions import FileOperationError


@dataclass
class DocumentConfig:
    """文档配置"""
    template_dir: str = "templates"
    output_dir: str = "outputs"
    enable_parallel: bool = True
    max_workers: int = 4
    enable_caching: bool = True
    cache_ttl: int = 3600
    enable_streaming: bool = True
    chunk_size: int = 1000
    auto_save: bool = True
    save_interval: float = 30.0


@dataclass
class DocumentTemplate:
    """文档模板"""
    name: str
    path: str
    variables: List[str] = field(default_factory=list)
    styles: Dict[str, Any] = field(default_factory=dict)
    sections: List[str] = field(default_factory=list)


class DocumentTemplateManager:
    """文档模板管理器"""
    
    def __init__(self, config: DocumentConfig):
        self.config = config
        self.templates: Dict[str, DocumentTemplate] = {}
        self.template_dir = Path(config.template_dir)
        self.template_dir.mkdir(exist_ok=True)
        
        # 加载模板
        self._load_templates()
    
    def _load_templates(self):
        """加载模板"""
        if not self.template_dir.exists():
            return
        
        for template_file in self.template_dir.glob("*.json"):
            try:
                with open(template_file, 'r', encoding='utf-8') as f:
                    template_data = json.load(f)
                
                template = DocumentTemplate(
                    name=template_data.get("name", template_file.stem),
                    path=str(template_file),
                    variables=template_data.get("variables", []),
                    styles=template_data.get("styles", {}),
                    sections=template_data.get("sections", [])
                )
                
                self.templates[template.name] = template
                
            except Exception as e:
                print(f"加载模板失败 {template_file}: {e}")
    
    def get_template(self, name: str) -> Optional[DocumentTemplate]:
        """获取模板"""
        return self.templates.get(name)
    
    def create_template(self, name: str, template_data: Dict[str, Any]) -> bool:
        """创建模板"""
        try:
            template_path = self.template_dir / f"{name}.json"
            with open(template_path, 'w', encoding='utf-8') as f:
                json.dump(template_data, f, ensure_ascii=False, indent=2)
            
            # 重新加载模板
            self._load_templates()
            return True
            
        except Exception as e:
            print(f"创建模板失败: {e}")
            return False
    
    def list_templates(self) -> List[str]:
        """列出所有模板"""
        return list(self.templates.keys())


class DocumentGenerator:
    """文档生成器"""
    
    def __init__(self, config: DocumentConfig):
        self.config = config
        self.template_manager = DocumentTemplateManager(config)
        self.metrics = get_metrics()
        self.cache = get_cache()
        self.executor = ThreadPoolExecutor(max_workers=config.max_workers)
        
        # 创建输出目录
        self.output_dir = Path(config.output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # 生成统计
        self.generation_stats = {
            "total_documents": 0,
            "successful_documents": 0,
            "failed_documents": 0,
            "average_generation_time": 0.0,
            "cache_hits": 0,
            "cache_misses": 0
        }
        self.stats_lock = threading.Lock()
    
    def generate_document(self, template_name: str, data: Dict[str, Any], 
                         output_path: str = None) -> str:
        """生成文档"""
        start_time = time.time()
        
        try:
            # 检查缓存
            if self.config.enable_caching:
                cache_key = self._generate_cache_key(template_name, data)
                cached_path = self.cache.get(cache_key)
                if cached_path and os.path.exists(cached_path):
                    with self.stats_lock:
                        self.generation_stats["cache_hits"] += 1
                    return cached_path
            
            # 生成文档
            if output_path is None:
                output_path = self._generate_output_path(template_name)
            
            # 获取模板
            template = self.template_manager.get_template(template_name)
            if not template:
                raise FileOperationError(f"模板不存在: {template_name}")
            
            # 创建文档
            doc = self._create_document_from_template(template, data)
            
            # 保存文档
            doc.save(output_path)
            
            # 缓存结果
            if self.config.enable_caching:
                cache_key = self._generate_cache_key(template_name, data)
                self.cache.set(cache_key, output_path, self.config.cache_ttl)
            
            # 更新统计
            generation_time = time.time() - start_time
            self._update_stats(True, generation_time)
            
            self.metrics.increment_counter("document_generation.successful")
            self.metrics.record_timer("document_generation.duration", generation_time)
            
            return output_path
            
        except Exception as e:
            generation_time = time.time() - start_time
            self._update_stats(False, generation_time)
            
            self.metrics.increment_counter("document_generation.failed")
            raise FileOperationError(f"文档生成失败: {e}")
    
    def batch_generate_documents(self, requests: List[Dict[str, Any]]) -> List[str]:
        """批量生成文档"""
        if not self.config.enable_parallel:
            # 串行处理
            results = []
            for request in requests:
                try:
                    result = self.generate_document(
                        request["template_name"],
                        request["data"],
                        request.get("output_path")
                    )
                    results.append(result)
                except Exception as e:
                    print(f"文档生成失败: {e}")
                    results.append(None)
            return results
        
        # 并行处理
        futures = []
        for request in requests:
            future = self.executor.submit(
                self.generate_document,
                request["template_name"],
                request["data"],
                request.get("output_path")
            )
            futures.append(future)
        
        results = []
        for future in as_completed(futures):
            try:
                result = future.result()
                results.append(result)
            except Exception as e:
                print(f"文档生成失败: {e}")
                results.append(None)
        
        return results
    
    def _create_document_from_template(self, template: DocumentTemplate, 
                                     data: Dict[str, Any]) -> Document:
        """从模板创建文档"""
        doc = Document()
        
        # 应用样式
        self._apply_styles(doc, template.styles)
        
        # 处理各个部分
        for section in template.sections:
            self._process_section(doc, section, data)
        
        return doc
    
    def _apply_styles(self, doc: Document, styles: Dict[str, Any]):
        """应用样式"""
        # 设置默认字体
        if "default_font" in styles:
            font_name = styles["default_font"]
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
        
        # 设置默认字号
        if "default_size" in styles:
            font_size = styles["default_size"]
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    
    def _process_section(self, doc: Document, section: str, data: Dict[str, Any]):
        """处理文档部分"""
        if section == "title":
            self._add_title(doc, data.get("title", ""))
        elif section == "content":
            self._add_content(doc, data.get("content", []))
        elif section == "table":
            self._add_table(doc, data.get("table", {}))
        elif section == "footer":
            self._add_footer(doc, data.get("footer", ""))
    
    def _add_title(self, doc: Document, title: str):
        """添加标题"""
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_content(self, doc: Document, content: List[Dict[str, Any]]):
        """添加内容"""
        for item in content:
            if item.get("type") == "heading":
                doc.add_heading(item.get("text", ""), item.get("level", 1))
            elif item.get("type") == "paragraph":
                doc.add_paragraph(item.get("text", ""))
            elif item.get("type") == "list":
                self._add_list(doc, item.get("items", []), item.get("ordered", False))
    
    def _add_list(self, doc: Document, items: List[str], ordered: bool = False):
        """添加列表"""
        for item in items:
            if ordered:
                doc.add_paragraph(item, style='List Number')
            else:
                doc.add_paragraph(item, style='List Bullet')
    
    def _add_table(self, doc: Document, table_data: Dict[str, Any]):
        """添加表格"""
        if "headers" not in table_data or "rows" not in table_data:
            return
        
        headers = table_data["headers"]
        rows = table_data["rows"]
        
        if not headers or not rows:
            return
        
        # 创建表格
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # 添加数据行
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_data in enumerate(row):
                if j < len(row_cells):
                    row_cells[j].text = str(cell_data)
    
    def _add_footer(self, doc: Document, footer: str):
        """添加页脚"""
        if footer:
            section = doc.sections[0]
            footer_para = section.footer.paragraphs[0]
            footer_para.text = footer
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _generate_cache_key(self, template_name: str, data: Dict[str, Any]) -> str:
        """生成缓存键"""
        import hashlib
        data_str = json.dumps(data, sort_keys=True)
        key_data = f"{template_name}:{data_str}"
        return hashlib.md5(key_data.encode()).hexdigest()
    
    def _generate_output_path(self, template_name: str) -> str:
        """生成输出路径"""
        timestamp = int(time.time())
        filename = f"{template_name}_{timestamp}.docx"
        return str(self.output_dir / filename)
    
    def _update_stats(self, success: bool, generation_time: float):
        """更新统计信息"""
        with self.stats_lock:
            self.generation_stats["total_documents"] += 1
            if success:
                self.generation_stats["successful_documents"] += 1
            else:
                self.generation_stats["failed_documents"] += 1
            
            # 更新平均生成时间
            total = self.generation_stats["total_documents"]
            current_avg = self.generation_stats["average_generation_time"]
            self.generation_stats["average_generation_time"] = (
                (current_avg * (total - 1) + generation_time) / total
            )
    
    def get_stats(self) -> Dict[str, Any]:
        """获取生成统计"""
        with self.stats_lock:
            total = self.generation_stats["total_documents"]
            success_rate = (
                self.generation_stats["successful_documents"] / max(1, total) * 100
            )
            
            return {
                **self.generation_stats,
                "success_rate": success_rate,
                "cache_hit_rate": (
                    self.generation_stats["cache_hits"] / 
                    max(1, self.generation_stats["cache_hits"] + self.generation_stats["cache_misses"]) * 100
                )
            }
    
    def cleanup(self):
        """清理资源"""
        if self.executor:
            self.executor.shutdown(wait=True)


# 全局文档生成器实例
_document_generator = None

def get_document_generator() -> DocumentGenerator:
    """获取文档生成器实例"""
    global _document_generator
    if _document_generator is None:
        _document_generator = DocumentGenerator(DocumentConfig())
    return _document_generator

