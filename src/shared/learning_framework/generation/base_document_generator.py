#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用文档生成框架
提供文档生成的通用接口和基础功能
"""

import os
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass
from datetime import datetime
from enum import Enum


class DocumentFormat(Enum):
    """文档格式枚举"""
    DOCX = "docx"      # Word文档
    PDF = "pdf"        # PDF文档
    HTML = "html"      # HTML文档
    TXT = "txt"        # 文本文档
    MD = "md"          # Markdown文档


class DocumentStyle(Enum):
    """文档样式枚举"""
    SIMPLE = "simple"          # 简单样式
    PROFESSIONAL = "professional"  # 专业样式
    COLORFUL = "colorful"      # 彩色样式
    MINIMAL = "minimal"        # 极简样式


@dataclass
class DocumentConfig:
    """文档配置"""
    title: str
    author: str = "Learning System"
    subject: str = "Learning Plan"
    output_format: DocumentFormat = DocumentFormat.DOCX
    style: DocumentStyle = DocumentStyle.PROFESSIONAL
    include_toc: bool = True
    include_page_numbers: bool = True
    include_headers: bool = True
    include_footers: bool = True
    font_size: int = 12
    line_spacing: float = 1.2
    margins: Dict[str, float] = None  # {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}
    
    def __post_init__(self):
        if self.margins is None:
            self.margins = {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}


@dataclass
class DocumentSection:
    """文档章节"""
    title: str
    content: Union[str, List[Dict[str, Any]]]
    level: int = 1  # 标题级别
    style: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


@dataclass
class DocumentTable:
    """文档表格"""
    headers: List[str]
    rows: List[List[str]]
    title: Optional[str] = None
    style: Optional[str] = None
    alignment: str = "left"


class BaseDocumentGenerator(ABC):
    """通用文档生成器基类"""
    
    def __init__(self, subject: str, config: Optional[Dict[str, Any]] = None):
        """
        初始化文档生成器
        
        Args:
            subject: 学科名称
            config: 配置参数
        """
        self.subject = subject
        self.config = config or {}
        self.templates: Dict[str, Dict[str, Any]] = {}
        self.styles: Dict[DocumentStyle, Dict[str, Any]] = {}
        self._init_templates()
        self._init_styles()
    
    @abstractmethod
    def _init_templates(self):
        """初始化学科特定的模板（抽象方法）"""
        pass
    
    @abstractmethod
    def _init_styles(self):
        """初始化学科特定的样式（抽象方法）"""
        pass
    
    def generate_document(self, sections: List[DocumentSection], 
                         config: DocumentConfig, output_path: Optional[str] = None) -> str:
        """
        生成文档
        
        Args:
            sections: 文档章节列表
            config: 文档配置
            output_path: 输出路径
            
        Returns:
            str: 生成的文档路径
        """
        if output_path is None:
            output_path = self._generate_output_path(config)
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # 根据格式生成文档
        if config.output_format == DocumentFormat.DOCX:
            return self._generate_docx(sections, config, output_path)
        elif config.output_format == DocumentFormat.HTML:
            return self._generate_html(sections, config, output_path)
        elif config.output_format == DocumentFormat.TXT:
            return self._generate_txt(sections, config, output_path)
        elif config.output_format == DocumentFormat.MD:
            return self._generate_markdown(sections, config, output_path)
        else:
            raise ValueError(f"不支持的文档格式: {config.output_format}")
    
    def _generate_output_path(self, config: DocumentConfig) -> str:
        """生成输出路径"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{config.title}_{timestamp}.{config.output_format.value}"
        return os.path.join(self._get_output_dir(), filename)
    
    def _get_output_dir(self) -> str:
        """获取输出目录"""
        return self.config.get('output_dir', 'outputs')
    
    def _generate_docx(self, sections: List[DocumentSection], 
                      config: DocumentConfig, output_path: str) -> str:
        """生成Word文档"""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
        
        doc = DocxDocument()
        self._setup_docx_styles(doc, config)
        
        # 添加标题
        doc.add_heading(config.title, 0)
        
        # 添加文档信息
        self._add_document_info(doc, config)
        
        # 添加目录
        if config.include_toc:
            self._add_table_of_contents(doc, sections)
        
        # 添加章节内容
        for section in sections:
            self._add_section_to_docx(doc, section, config)
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def _generate_html(self, sections: List[DocumentSection], 
                      config: DocumentConfig, output_path: str) -> str:
        """生成HTML文档"""
        html_content = self._generate_html_content(sections, config)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def _generate_txt(self, sections: List[DocumentSection], 
                     config: DocumentConfig, output_path: str) -> str:
        """生成文本文档"""
        txt_content = self._generate_txt_content(sections, config)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(txt_content)
        
        return output_path
    
    def _generate_markdown(self, sections: List[DocumentSection], 
                          config: DocumentConfig, output_path: str) -> str:
        """生成Markdown文档"""
        md_content = self._generate_markdown_content(sections, config)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return output_path
    
    def _setup_docx_styles(self, doc, config: DocumentConfig):
        """设置Word文档样式"""
        # 设置字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(config.font_size)
        
        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = config.line_spacing
        
        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(config.margins['top'])
            section.bottom_margin = Inches(config.margins['bottom'])
            section.left_margin = Inches(config.margins['left'])
            section.right_margin = Inches(config.margins['right'])
    
    def _add_document_info(self, doc, config: DocumentConfig):
        """添加文档信息"""
        info_para = doc.add_paragraph()
        info_para.add_run(f"作者: {config.author}\n")
        info_para.add_run(f"学科: {config.subject}\n")
        info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.paragraph_format.line_spacing = 1.2
        
        # 添加分页符
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc, sections: List[DocumentSection]):
        """添加目录"""
        doc.add_heading("目录", 1)
        
        toc_para = doc.add_paragraph()
        for i, section in enumerate(sections, 1):
            indent = "  " * (section.level - 1)
            toc_para.add_run(f"{indent}{i}. {section.title}\n")
        
        doc.add_page_break()
    
    def _add_section_to_docx(self, doc, section: DocumentSection, config: DocumentConfig):
        """添加章节到Word文档"""
        # 添加标题
        doc.add_heading(section.title, section.level)
        
        # 添加内容
        if isinstance(section.content, str):
            para = doc.add_paragraph(section.content)
            para.paragraph_format.line_spacing = config.line_spacing
        elif isinstance(section.content, list):
            for item in section.content:
                if isinstance(item, dict):
                    if item.get('type') == 'table':
                        self._add_table_to_docx(doc, item, config)
                    elif item.get('type') == 'list':
                        self._add_list_to_docx(doc, item, config)
                    else:
                        para = doc.add_paragraph(str(item.get('content', '')))
                        para.paragraph_format.line_spacing = config.line_spacing
                else:
                    para = doc.add_paragraph(str(item))
                    para.paragraph_format.line_spacing = config.line_spacing
    
    def _add_table_to_docx(self, doc, table_data: Dict[str, Any], config: DocumentConfig):
        """添加表格到Word文档"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return
        
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # 添加数据行
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_data in enumerate(row):
                if j < len(row_cells):
                    row_cells[j].text = str(cell_data)
    
    def _add_list_to_docx(self, doc, list_data: Dict[str, Any], config: DocumentConfig):
        """添加列表到Word文档"""
        items = list_data.get('items', [])
        list_type = list_data.get('style', 'bullet')
        
        for item in items:
            para = doc.add_paragraph(str(item), style='List Bullet' if list_type == 'bullet' else 'List Number')
            para.paragraph_format.line_spacing = config.line_spacing
    
    def _generate_html_content(self, sections: List[DocumentSection], config: DocumentConfig) -> str:
        """生成HTML内容"""
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{config.title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: {config.line_spacing}; }}
        h1, h2, h3, h4, h5, h6 {{ color: #333; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        .toc {{ background-color: #f9f9f9; padding: 20px; margin: 20px 0; }}
    </style>
</head>
<body>
    <h1>{config.title}</h1>
    <div class="toc">
        <h2>目录</h2>
        <ul>
"""
        
        for i, section in enumerate(sections, 1):
            html += f"            <li><a href=\"#section{i}\">{section.title}</a></li>\n"
        
        html += "        </ul>\n    </div>\n"
        
        for i, section in enumerate(sections, 1):
            html += f"    <h{section.level} id=\"section{i}\">{section.title}</h{section.level}>\n"
            
            if isinstance(section.content, str):
                html += f"    <p>{section.content}</p>\n"
            elif isinstance(section.content, list):
                for item in section.content:
                    if isinstance(item, dict):
                        if item.get('type') == 'table':
                            html += self._generate_html_table(item)
                        elif item.get('type') == 'list':
                            html += self._generate_html_list(item)
                        else:
                            html += f"    <p>{item.get('content', '')}</p>\n"
                    else:
                        html += f"    <p>{item}</p>\n"
        
        html += "</body>\n</html>"
        return html
    
    def _generate_html_table(self, table_data: Dict[str, Any]) -> str:
        """生成HTML表格"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return ""
        
        html = "    <table>\n        <thead>\n            <tr>\n"
        for header in headers:
            html += f"                <th>{header}</th>\n"
        html += "            </tr>\n        </thead>\n        <tbody>\n"
        
        for row in rows:
            html += "            <tr>\n"
            for cell in row:
                html += f"                <td>{cell}</td>\n"
            html += "            </tr>\n"
        
        html += "        </tbody>\n    </table>\n"
        return html
    
    def _generate_html_list(self, list_data: Dict[str, Any]) -> str:
        """生成HTML列表"""
        items = list_data.get('items', [])
        list_type = list_data.get('style', 'bullet')
        
        if not items:
            return ""
        
        tag = "ul" if list_type == "bullet" else "ol"
        html = f"    <{tag}>\n"
        for item in items:
            html += f"        <li>{item}</li>\n"
        html += f"    </{tag}>\n"
        return html
    
    def _generate_txt_content(self, sections: List[DocumentSection], config: DocumentConfig) -> str:
        """生成文本内容"""
        content = f"{config.title}\n"
        content += "=" * len(config.title) + "\n\n"
        content += f"作者: {config.author}\n"
        content += f"学科: {config.subject}\n"
        content += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        for section in sections:
            content += f"{'#' * section.level} {section.title}\n\n"
            
            if isinstance(section.content, str):
                content += f"{section.content}\n\n"
            elif isinstance(section.content, list):
                for item in section.content:
                    if isinstance(item, dict):
                        if item.get('type') == 'table':
                            content += self._generate_txt_table(item)
                        elif item.get('type') == 'list':
                            content += self._generate_txt_list(item)
                        else:
                            content += f"{item.get('content', '')}\n\n"
                    else:
                        content += f"{item}\n\n"
        
        return content
    
    def _generate_txt_table(self, table_data: Dict[str, Any]) -> str:
        """生成文本表格"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return ""
        
        # 计算列宽
        col_widths = [len(header) for header in headers]
        for row in rows:
            for i, cell in enumerate(row):
                if i < len(col_widths):
                    col_widths[i] = max(col_widths[i], len(str(cell)))
        
        # 生成表格
        content = ""
        # 表头
        content += "| " + " | ".join(header.ljust(col_widths[i]) for i, header in enumerate(headers)) + " |\n"
        content += "|" + "|".join("-" * (width + 2) for width in col_widths) + "|\n"
        
        # 数据行
        for row in rows:
            content += "| " + " | ".join(str(cell).ljust(col_widths[i]) for i, cell in enumerate(row)) + " |\n"
        
        content += "\n"
        return content
    
    def _generate_txt_list(self, list_data: Dict[str, Any]) -> str:
        """生成文本列表"""
        items = list_data.get('items', [])
        list_type = list_data.get('style', 'bullet')
        
        if not items:
            return ""
        
        content = ""
        for i, item in enumerate(items, 1):
            if list_type == "bullet":
                content += f"• {item}\n"
            else:
                content += f"{i}. {item}\n"
        
        content += "\n"
        return content
    
    def _generate_markdown_content(self, sections: List[DocumentSection], config: DocumentConfig) -> str:
        """生成Markdown内容"""
        content = f"# {config.title}\n\n"
        content += f"**作者**: {config.author}  \n"
        content += f"**学科**: {config.subject}  \n"
        content += f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n\n"
        
        for section in sections:
            content += f"{'#' * (section.level + 1)} {section.title}\n\n"
            
            if isinstance(section.content, str):
                content += f"{section.content}\n\n"
            elif isinstance(section.content, list):
                for item in section.content:
                    if isinstance(item, dict):
                        if item.get('type') == 'table':
                            content += self._generate_markdown_table(item)
                        elif item.get('type') == 'list':
                            content += self._generate_markdown_list(item)
                        else:
                            content += f"{item.get('content', '')}\n\n"
                    else:
                        content += f"{item}\n\n"
        
        return content
    
    def _generate_markdown_table(self, table_data: Dict[str, Any]) -> str:
        """生成Markdown表格"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return ""
        
        content = "| " + " | ".join(headers) + " |\n"
        content += "| " + " | ".join("---" for _ in headers) + " |\n"
        
        for row in rows:
            content += "| " + " | ".join(str(cell) for cell in row) + " |\n"
        
        content += "\n"
        return content
    
    def _generate_markdown_list(self, list_data: Dict[str, Any]) -> str:
        """生成Markdown列表"""
        items = list_data.get('items', [])
        list_type = list_data.get('style', 'bullet')
        
        if not items:
            return ""
        
        content = ""
        for i, item in enumerate(items, 1):
            if list_type == "bullet":
                content += f"- {item}\n"
            else:
                content += f"{i}. {item}\n"
        
        content += "\n"
        return content
    
    def add_template(self, template_name: str, template: Dict[str, Any]):
        """添加模板"""
        self.templates[template_name] = template
    
    def get_template(self, template_name: str) -> Optional[Dict[str, Any]]:
        """获取模板"""
        return self.templates.get(template_name)
    
    def list_templates(self) -> List[str]:
        """列出所有模板"""
        return list(self.templates.keys())
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        """获取支持的文档格式"""
        return list(DocumentFormat)
    
    def get_supported_styles(self) -> List[DocumentStyle]:
        """获取支持的文档样式"""
        return list(DocumentStyle)
