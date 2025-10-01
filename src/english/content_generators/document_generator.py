#!/usr/bin/env python3
"""
文档生成器 - 专门负责Word文档的生成和格式化
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent.parent.parent
sys.path.append(str(project_root))


class DocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.project_root = Path(__file__).parent.parent.parent.parent
    
    def _set_paragraph_spacing(self, paragraph):
        """设置段落行间距为单倍行距"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 设置为单倍行距

    def generate_word_document(self, content: Dict, filename: str = None) -> str:
        """生成Word文档"""
        try:
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('英语学习内容', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 根据内容类型添加相应部分
            if 'day' in content:
                # 单日内容
                self._add_single_day_to_document(doc, content)
            elif 'days' in content:
                # 多日内容
                for day_content in content['days']:
                    self._add_single_day_to_document(doc, day_content)
                    doc.add_page_break()
            
            # 保存文档
            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                if 'day' in content:
                    plan_id = content.get('plan_id', 'unknown')
                    day_num = content.get('day', 1)
                    date_str = datetime.now().strftime("%m%d_%H%M")
                    filename = f"day{day_num}_{date_str}.docx"
                else:
                    filename = f"learning_content_{timestamp}.docx"
            
            # 确保文件名以.docx结尾
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            # 根据计划ID创建目录结构
            plan_id = content.get('plan_id', 'default')
            output_dir = self.project_root / "outputs" / "english" / "word_documents" / plan_id
            output_dir.mkdir(parents=True, exist_ok=True)
            
            output_path = output_dir / filename
            doc.save(str(output_path))
            
            print(f"📄 Word文档已保存到: {output_path}")
            return str(output_path)
            
        except Exception as e:
            print(f"❌ 生成Word文档失败: {e}")
            return ""

    def _add_single_day_to_document(self, doc: Document, content: Dict) -> None:
        """向文档添加单日内容"""
        # 添加日期标题
        day_title = f"第{content.get('day', 1)}天 - {content.get('date', datetime.now().strftime('%Y-%m-%d'))}"
        doc.add_heading(day_title, level=1)
        
        # 添加各个部分
        if 'vocabulary' in content:
            self._add_vocabulary_section(doc, content['vocabulary'])
        
        if 'morphology' in content:
            self._add_morphology_section(doc, content['morphology'])
        
        if 'syntax' in content:
            self._add_syntax_section(doc, content['syntax'])
        
        if 'practice' in content:
            self._add_practice_section(doc, content['practice'])

    def _add_vocabulary_section(self, doc: Document, vocabulary: Dict) -> None:
        """添加词汇部分"""
        doc.add_heading('📚 词汇学习', level=2)
        
        # 新学单词
        if 'new_words' in vocabulary:
            doc.add_heading('🆕 新学单词', level=3)
            
            for category, words in vocabulary['new_words'].items():
                if words:
                    category_name = {
                        'core_functional': '核心功能词',
                        'connectors_relational': '连接关系词', 
                        'auxiliary_supplemental': '辅助补充词'
                    }.get(category, category)
                    
                    doc.add_heading(f'{category_name}', level=4)
                    
                    # 创建表格
                    table = doc.add_table(rows=len(words)+1, cols=3)
                    table.style = 'Table Grid'
                    
                    # 表头
                    header_cells = table.rows[0].cells
                    header_cells[0].text = '单词'
                    header_cells[1].text = '词性'
                    header_cells[2].text = '定义'
                    
                    # 填充数据
                    for i, word in enumerate(words, 1):
                        row_cells = table.rows[i].cells
                        row_cells[0].text = word.get('word', '')
                        row_cells[1].text = word.get('part_of_speech', '')
                        row_cells[2].text = word.get('definition', '')
        
        # 复习单词
        if 'review_words' in vocabulary:
            review_words = vocabulary['review_words']
            if review_words:
                doc.add_heading('🔄 复习单词', level=3)
                
                # 创建表格
                table = doc.add_table(rows=len(review_words)+1, cols=3)
                table.style = 'Table Grid'
                
                # 表头
                header_cells = table.rows[0].cells
                header_cells[0].text = '单词'
                header_cells[1].text = '词性'
                header_cells[2].text = '定义'
                
                # 填充数据
                for i, word in enumerate(review_words, 1):
                    row_cells = table.rows[i].cells
                    if isinstance(word, dict):
                        row_cells[0].text = word.get('word', '')
                        row_cells[1].text = word.get('part_of_speech', '')
                        row_cells[2].text = word.get('definition', '')
                    else:
                        row_cells[0].text = str(word)
                        row_cells[1].text = ''
                        row_cells[2].text = ''

    def _add_morphology_section(self, doc: Document, morphology: Dict) -> None:
        """添加词法部分"""
        doc.add_heading('🔤 词法学习', level=2)
        
        learning_points = morphology.get('learning_points', [])
        if learning_points:
            for point in learning_points:
                doc.add_heading(f"{point.get('name', '词法规则')}", level=3)
                type_p = doc.add_paragraph(f"类型: {point.get('type', 'N/A')}")
                self._set_paragraph_spacing(type_p)
                desc_p = doc.add_paragraph(f"描述: {point.get('description', 'N/A')}")
                self._set_paragraph_spacing(desc_p)
                
                if point.get('rules'):
                    rules_p = doc.add_paragraph(f"规则: {point['rules']}")
                    self._set_paragraph_spacing(rules_p)
                
                if point.get('examples'):
                    examples_p = doc.add_paragraph("例句:")
                    self._set_paragraph_spacing(examples_p)
                    for example in point['examples'][:3]:  # 最多显示3个例句
                        example_p = doc.add_paragraph(f"  • {example}")
                        self._set_paragraph_spacing(example_p)

    def _add_syntax_section(self, doc: Document, syntax: Dict) -> None:
        """添加句法部分"""
        doc.add_heading('📝 句法学习', level=2)
        
        learning_points = syntax.get('learning_points', [])
        if learning_points:
            for point in learning_points:
                doc.add_heading(f"{point.get('name', '句法规则')}", level=3)
                type_p = doc.add_paragraph(f"类型: {point.get('type', 'N/A')}")
                self._set_paragraph_spacing(type_p)
                struct_p = doc.add_paragraph(f"结构: {point.get('structure', 'N/A')}")
                self._set_paragraph_spacing(struct_p)
                
                if point.get('description'):
                    desc_p = doc.add_paragraph(f"描述: {point['description']}")
                    self._set_paragraph_spacing(desc_p)
                
                if point.get('examples'):
                    examples_p = doc.add_paragraph("例句:")
                    self._set_paragraph_spacing(examples_p)
                    for example in point['examples'][:3]:  # 最多显示3个例句
                        example_p = doc.add_paragraph(f"  • {example}")
                        self._set_paragraph_spacing(example_p)

    def _add_practice_section(self, doc: Document, practice: Dict) -> None:
        """添加练习部分"""
        doc.add_heading('💪 练习部分', level=2)
        
        # 收集练习题答案，稍后单独输出
        exercise_answers = []
        
        # 练习句子
        if 'practice_sentences' in practice:
            sentences_data = practice['practice_sentences']
            if isinstance(sentences_data, dict) and 'practice_sentences' in sentences_data:
                sentences = sentences_data['practice_sentences']
                if sentences:
                    doc.add_heading('📝 练习句子', level=3)
                    for i, sentence in enumerate(sentences, 1):
                        p1 = doc.add_paragraph(f"{i}. {sentence.get('sentence', '')}")
                        self._set_paragraph_spacing(p1)
                        p2 = doc.add_paragraph(f"   翻译: {sentence.get('translation', '')}")
                        self._set_paragraph_spacing(p2)
                        p3 = doc.add_paragraph(f"   词法规则: {sentence.get('morphology_rule', '')}")
                        self._set_paragraph_spacing(p3)
                        p4 = doc.add_paragraph(f"   句法结构: {sentence.get('syntactic_structure', '')}")
                        self._set_paragraph_spacing(p4)
                        p5 = doc.add_paragraph(f"   解析: {sentence.get('explanation', '')}")
                        self._set_paragraph_spacing(p5)
                        # 去掉空行分隔
        
        # 练习题（不显示答案）
        if 'practice_exercises' in practice:
            exercises_data = practice['practice_exercises']
            if isinstance(exercises_data, dict) and 'practice_exercises' in exercises_data:
                exercises = exercises_data['practice_exercises']
                if exercises:
                    doc.add_heading('📋 练习题', level=3)
                    for exercise in exercises:
                        ex_id = exercise.get('id', '')
                        ex_type = exercise.get('type', '')
                        
                        # 去掉题型标记[]
                        question_p = doc.add_paragraph(f"{ex_id}. {exercise.get('question', '')}")
                        self._set_paragraph_spacing(question_p)
                        
                        # 收集答案信息
                        answer_info = {
                            'id': ex_id,
                            'type': ex_type,
                            'question': exercise.get('question', ''),
                            'answer': '',
                            'explanation': exercise.get('explanation', '')
                        }
                        
                        if ex_type == 'choice' and 'options' in exercise:
                            # 为选择题添加A、B、C、D标题
                            option_labels = ['A', 'B', 'C', 'D']
                            for i, option in enumerate(exercise['options']):
                                if i < len(option_labels):
                                    option_p = doc.add_paragraph(f"   {option_labels[i]}. {option}")
                                    self._set_paragraph_spacing(option_p)
                            # 收集答案，不在题目中显示
                            answer_info['answer'] = exercise.get('correct_answer', '')
                        
                        elif ex_type == 'translation':
                            # 翻译题只显示题目内容，不显示答案
                            # 根据题目描述判断翻译方向
                            question = exercise.get('question', '').lower()
                            if 'chinese_text' in exercise and 'english_text' in exercise:
                                if '中文翻译成英文' in exercise.get('question', '') or 'chinese to english' in question:
                                    # 中译英：显示中文，答案是英文
                                    chinese_p = doc.add_paragraph(f"   中文: {exercise['chinese_text']}")
                                    self._set_paragraph_spacing(chinese_p)
                                    answer_info['answer'] = exercise.get('english_text', '')
                                elif '英文翻译成中文' in exercise.get('question', '') or 'english to chinese' in question:
                                    # 英译中：显示英文，答案是中文
                                    english_p = doc.add_paragraph(f"   英文: {exercise['english_text']}")
                                    self._set_paragraph_spacing(english_p)
                                    answer_info['answer'] = exercise.get('chinese_text', '')
                                else:
                                    # 默认中译英
                                    chinese_p = doc.add_paragraph(f"   中文: {exercise['chinese_text']}")
                                    self._set_paragraph_spacing(chinese_p)
                                    answer_info['answer'] = exercise.get('english_text', '')
                            elif 'chinese_text' in exercise:
                                chinese_p = doc.add_paragraph(f"   中文: {exercise['chinese_text']}")
                                self._set_paragraph_spacing(chinese_p)
                                answer_info['answer'] = exercise.get('english_text', '')
                            elif 'english_text' in exercise:
                                english_p = doc.add_paragraph(f"   英文: {exercise['english_text']}")
                                self._set_paragraph_spacing(english_p)
                                answer_info['answer'] = exercise.get('chinese_text', '')
                        
                        elif ex_type == 'fill_blank':
                            if 'sentence' in exercise:
                                sentence_p = doc.add_paragraph(f"   句子: {exercise['sentence']}")
                                self._set_paragraph_spacing(sentence_p)
                            # 收集答案，不在题目中显示
                            answer_info['answer'] = exercise.get('answer', '')
                        
                        exercise_answers.append(answer_info)
                        # 去掉空行分隔
        
        # 添加答案页（如果有练习题）
        if exercise_answers:
            self._add_answers_page(doc, exercise_answers)

    def _add_answers_page(self, doc: Document, exercise_answers: List[Dict]) -> None:
        """添加答案页"""
        # 添加分页符
        doc.add_page_break()
        
        # 添加答案页标题
        doc.add_heading('📋 练习题答案', level=2)
        
        for answer_info in exercise_answers:
            ex_id = answer_info['id']
            ex_type = answer_info['type']
            answer = answer_info['answer']
            explanation = answer_info['explanation']
            
            # 题目编号（去掉题型标记[]）
            id_p = doc.add_paragraph(f"{ex_id}.")
            self._set_paragraph_spacing(id_p)
            
            # 答案
            if answer:
                answer_p = doc.add_paragraph(f"   答案: {answer}")
                self._set_paragraph_spacing(answer_p)
            
            # 解析
            if explanation:
                explanation_p = doc.add_paragraph(f"   解析: {explanation}")
                self._set_paragraph_spacing(explanation_p)
            
            # 去掉空行分隔
