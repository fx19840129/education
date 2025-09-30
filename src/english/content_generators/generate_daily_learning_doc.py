#!/usr/bin/env python3
"""
每日学习内容文档生成器
生成包含单词、词法、句法、练习句子、练习题的完整学习文档
"""

import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent.parent.parent
sys.path.append(str(project_root))

from src.english.learning_content_generator import LearningContentGenerator
from src.english.services.fsrs_learning_generator import FSRSLearningGenerator
from src.english.generate_morphology_content import MorphologyContentGenerator
from src.english.generate_syntax_content import SyntaxContentGenerator
from src.english.generate_practice_sentences import PracticeSentencesGenerator
from src.english.generate_practice_exercises import PracticeExercisesGenerator

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx 未安装，将生成Markdown格式文档")


class DailyLearningDocumentGenerator:
    """每日学习内容文档生成器"""
    
    def __init__(self):
        self.plan_reader = LearningContentGenerator()
        self.fsrs_generator = FSRSLearningGenerator()
        self.morphology_generator = MorphologyContentGenerator()
        self.syntax_generator = SyntaxContentGenerator()
        self.sentence_generator = PracticeSentencesGenerator()
        self.exercise_generator = PracticeExercisesGenerator()
        
        # 词性中文名称映射
        self.pos_chinese_names = {
            'noun': '名词',
            'verb': '动词',
            'adjective': '形容词',
            'adverb': '副词',
            'preposition': '介词',
            'pronoun': '代词',
            'conjunction': '连词',
            'article': '冠词',
            'determiner': '限定词',
            'interjection': '感叹词',
            'numeral': '数词',
            'modal': '情态动词',
            'phrase': '短语',
            'auxiliary': '助动词'
        }
    
    def generate_daily_learning_document(self, target_date: str = None, plan_id: str = None) -> Dict:
        """生成指定日期的完整学习文档"""
        if target_date is None:
            target_date = datetime.now().strftime("%Y-%m-%d")
        
        # 获取学习计划
        if plan_id is None:
            plans = self.plan_reader.list_plans()
            if not plans:
                raise ValueError("没有找到学习计划")
            plan_id = plans[0]['id']
        
        learning_plan = self.plan_reader.read_plan(plan_id)
        if not learning_plan:
            raise ValueError(f"无法读取学习计划: {plan_id}")
        
        print(f"📅 生成 {target_date} 的完整学习内容...")
        
        # 生成各种学习内容
        daily_words = self.fsrs_generator.generate_daily_learning_content(learning_plan, target_date)
        daily_morphology = self.morphology_generator.generate_daily_morphology(learning_plan, target_date)
        daily_syntax = self.syntax_generator.generate_daily_syntax(learning_plan, target_date)
        daily_sentences = self.sentence_generator.generate_daily_sentences(learning_plan, target_date)
        daily_exercises = self.exercise_generator.generate_daily_exercises(learning_plan, target_date)
        
        # 整合所有内容
        daily_content = {
            "date": target_date,
            "stage": learning_plan.get("metadata", {}).get("stage", ""),
            "plan_name": learning_plan.get("metadata", {}).get("plan_name", ""),
            "words": daily_words,
            "morphology": daily_morphology,
            "syntax": daily_syntax,
            "sentences": daily_sentences,
            "exercises": daily_exercises
        }
        
        return daily_content
    
    def generate_word_document(self, daily_content: Dict, output_path: str = None) -> str:
        """生成Word文档"""
        if not DOCX_AVAILABLE:
            return self.generate_markdown_document(daily_content, output_path)
        
        if output_path is None:
            output_path = f"outputs/english/daily_learning_{daily_content['date']}.docx"
        
        # 确保输出目录存在
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        self._setup_document_styles(doc)
        
        # 添加标题
        title = doc.add_heading(f"英语学习内容 - {daily_content['date']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        self._add_basic_info(doc, daily_content)
        
        # 添加单词学习内容
        self._add_words_content(doc, daily_content['words'])
        
        # 添加词法学习内容
        self._add_morphology_content(doc, daily_content['morphology'])
        
        # 添加句法学习内容
        self._add_syntax_content(doc, daily_content['syntax'])
        
        # 添加练习句子
        self._add_sentences_content(doc, daily_content['sentences'])
        
        # 添加练习题
        self._add_exercises_content(doc, daily_content['exercises'])
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def generate_markdown_document(self, daily_content: Dict, output_path: str = None) -> str:
        """生成Markdown文档"""
        if output_path is None:
            output_path = f"outputs/english/daily_learning_{daily_content['date']}.md"
        
        # 确保输出目录存在
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        content = []
        
        # 标题
        content.append(f"# 英语学习内容 - {daily_content['date']}")
        content.append("")
        
        # 基本信息
        content.append("## 基本信息")
        content.append(f"- **学习阶段**: {daily_content['stage']}")
        content.append(f"- **学习计划**: {daily_content['plan_name']}")
        content.append(f"- **学习日期**: {daily_content['date']}")
        content.append("")
        
        # 单词学习内容
        content.append("## 📚 单词学习")
        self._add_words_markdown(content, daily_content['words'])
        
        # 词法学习内容
        content.append("## 📖 词法学习")
        self._add_morphology_markdown(content, daily_content['morphology'])
        
        # 句法学习内容
        content.append("## 📝 句法学习")
        self._add_syntax_markdown(content, daily_content['syntax'])
        
        # 练习句子
        content.append("## 💬 练习句子")
        self._add_sentences_markdown(content, daily_content['sentences'])
        
        # 练习题
        content.append("## ✏️ 练习题")
        self._add_exercises_markdown(content, daily_content['exercises'])
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        return output_path
    
    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置标题样式
        styles = doc.styles
        
        # 创建自定义样式
        if 'Custom Heading 2' not in [style.name for style in styles]:
            heading2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.paragraph_format.space_before = Pt(12)
            heading2_style.paragraph_format.space_after = Pt(6)
    
    def _add_basic_info(self, doc, daily_content):
        """添加基本信息"""
        doc.add_heading("基本信息", level=2)
        
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = "学习阶段"
        info_table.cell(0, 1).text = daily_content['stage']
        info_table.cell(1, 0).text = "学习计划"
        info_table.cell(1, 1).text = daily_content['plan_name']
        info_table.cell(2, 0).text = "学习日期"
        info_table.cell(2, 1).text = daily_content['date']
        
        doc.add_paragraph()
    
    def _add_words_content(self, doc, words_content):
        """添加单词学习内容"""
        doc.add_heading("📚 单词学习", level=2)
        
        if words_content.get('pos_content'):
            for pos, words in words_content['pos_content'].items():
                # 获取词性的中文名称
                chinese_name = self.pos_chinese_names.get(pos, pos)
                doc.add_heading(f"{chinese_name} ({pos}) - {len(words)}个", level=3)
                
                # 创建单词表格
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # 设置表头
                header_cells = table.rows[0].cells
                header_cells[0].text = "序号"
                header_cells[1].text = "单词"
                header_cells[2].text = "中文解释"
                header_cells[3].text = "难度"
                
                # 设置表头样式
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加单词数据
                for i, word in enumerate(words, 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = word['word']
                    row_cells[2].text = word.get('translation', '')
                    row_cells[3].text = f"{word.get('difficulty', 0):.1f}" if word.get('difficulty') else "0.0"
                    
                    # 设置单元格对齐方式
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置表格列宽
                table.columns[0].width = Inches(0.5)  # 序号列
                table.columns[1].width = Inches(1.2)  # 单词列
                table.columns[2].width = Inches(2.0)  # 中文解释列
                table.columns[3].width = Inches(0.8)  # 难度列
                
                doc.add_paragraph()  # 添加空行
    
    def _add_morphology_content(self, doc, morphology_content):
        """添加词法学习内容"""
        doc.add_heading("📖 词法学习", level=2)
        
        if morphology_content.get('morphology_items'):
            for i, item in enumerate(morphology_content['morphology_items'], 1):
                doc.add_heading(f"{i}. {item['name']}", level=3)
                
                p = doc.add_paragraph()
                p.add_run("类型: ").bold = True
                p.add_run(item['type'])
                
                p = doc.add_paragraph()
                p.add_run("描述: ").bold = True
                p.add_run(item['description'])
                
                if item.get('rules'):
                    p = doc.add_paragraph()
                    p.add_run("规则: ").bold = True
                    for rule in item['rules']:
                        doc.add_paragraph(f"• {rule}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_syntax_content(self, doc, syntax_content):
        """添加句法学习内容"""
        doc.add_heading("📝 句法学习", level=2)
        
        if syntax_content.get('syntax_items'):
            for i, item in enumerate(syntax_content['syntax_items'], 1):
                doc.add_heading(f"{i}. {item['name']}", level=3)
                
                p = doc.add_paragraph()
                p.add_run("类型: ").bold = True
                p.add_run(item['type'])
                
                p = doc.add_paragraph()
                p.add_run("结构: ").bold = True
                p.add_run(item['structure'])
                
                if item.get('examples'):
                    p = doc.add_paragraph()
                    p.add_run("例句: ").bold = True
                    for example in item['examples']:
                        doc.add_paragraph(f"• {example}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_sentences_content(self, doc, sentences_content):
        """添加练习句子内容"""
        doc.add_heading("💬 练习句子", level=2)
        
        if sentences_content.get('practice_sentences'):
            for i, sentence in enumerate(sentences_content['practice_sentences'], 1):
                doc.add_heading(f"句子 {i}", level=3)
                
                p = doc.add_paragraph()
                p.add_run("英文: ").bold = True
                p.add_run(sentence['sentence'])
                
                p = doc.add_paragraph()
                p.add_run("中文: ").bold = True
                p.add_run(sentence['translation'])
                
                if sentence.get('target_words'):
                    p = doc.add_paragraph()
                    p.add_run("目标单词: ").bold = True
                    p.add_run(", ".join(sentence['target_words']))
                
                if sentence.get('morphology_points'):
                    p = doc.add_paragraph()
                    p.add_run("词法点: ").bold = True
                    p.add_run(", ".join(sentence['morphology_points']))
                
                if sentence.get('syntax_structure'):
                    p = doc.add_paragraph()
                    p.add_run("句法结构: ").bold = True
                    p.add_run(sentence['syntax_structure'])
        
        doc.add_paragraph()
    
    def _add_exercises_content(self, doc, exercises_content):
        """添加练习题内容"""
        doc.add_heading("✏️ 练习题", level=2)
        
        if exercises_content.get('practice_exercises'):
            # 按题型分组
            type_groups = {}
            for exercise in exercises_content['practice_exercises']:
                ex_type = exercise['type']
                if ex_type not in type_groups:
                    type_groups[ex_type] = []
                type_groups[ex_type].append(exercise)
            
            for ex_type, exercises in type_groups.items():
                type_name = {"choice": "选择题", "translation": "翻译题", "fill_blank": "填空题"}.get(ex_type, ex_type)
                doc.add_heading(f"{type_name} ({len(exercises)}个)", level=3)
                
                for i, exercise in enumerate(exercises, 1):
                    doc.add_heading(f"{i}. {exercise['question']}", level=4)
                    
                    if exercise['type'] == 'choice':
                        for j, option in enumerate(exercise['options']):
                            doc.add_paragraph(f"{chr(65+j)}. {option}")
                        p = doc.add_paragraph()
                        p.add_run("答案: ").bold = True
                        p.add_run(exercise['correct_answer'])
                    
                    elif exercise['type'] == 'translation':
                        p = doc.add_paragraph()
                        p.add_run("中文: ").bold = True
                        p.add_run(exercise['chinese_text'])
                        p = doc.add_paragraph()
                        p.add_run("英文: ").bold = True
                        p.add_run(exercise['english_text'])
                    
                    elif exercise['type'] == 'fill_blank':
                        p = doc.add_paragraph()
                        p.add_run("句子: ").bold = True
                        p.add_run(exercise['sentence'])
                        p = doc.add_paragraph()
                        p.add_run("答案: ").bold = True
                        p.add_run(exercise['answer'])
                    
                    p = doc.add_paragraph()
                    p.add_run("解析: ").bold = True
                    p.add_run(exercise['explanation'])
                    
                    doc.add_paragraph()
    
    def _add_words_markdown(self, content, words_content):
        """添加单词Markdown内容"""
        if words_content.get('pos_content'):
            for pos, words in words_content['pos_content'].items():
                # 获取词性的中文名称
                chinese_name = self.pos_chinese_names.get(pos, pos)
                content.append(f"### {chinese_name} ({pos}) - {len(words)}个")
                content.append("")
                
                # 创建Markdown表格
                content.append("| 序号 | 单词 | 中文解释 | 难度 |")
                content.append("|------|------|----------|------|")
                
                for i, word in enumerate(words, 1):
                    difficulty = f"{word.get('difficulty', 0):.1f}" if word.get('difficulty') else "0.0"
                    content.append(f"| {i} | **{word['word']}** | {word.get('translation', '')} | {difficulty} |")
                
                content.append("")
    
    def _add_morphology_markdown(self, content, morphology_content):
        """添加词法Markdown内容"""
        if morphology_content.get('morphology_items'):
            for i, item in enumerate(morphology_content['morphology_items'], 1):
                content.append(f"### {i}. {item['name']}")
                content.append("")
                content.append(f"**类型**: {item['type']}")
                content.append(f"**描述**: {item['description']}")
                if item.get('rules'):
                    content.append("**规则**:")
                    for rule in item['rules']:
                        content.append(f"- {rule}")
                content.append("")
    
    def _add_syntax_markdown(self, content, syntax_content):
        """添加句法Markdown内容"""
        if syntax_content.get('syntax_items'):
            for i, item in enumerate(syntax_content['syntax_items'], 1):
                content.append(f"### {i}. {item['name']}")
                content.append("")
                content.append(f"**类型**: {item['type']}")
                content.append(f"**结构**: {item['structure']}")
                if item.get('examples'):
                    content.append("**例句**:")
                    for example in item['examples']:
                        content.append(f"- {example}")
                content.append("")
    
    def _add_sentences_markdown(self, content, sentences_content):
        """添加练习句子Markdown内容"""
        if sentences_content.get('practice_sentences'):
            for i, sentence in enumerate(sentences_content['practice_sentences'], 1):
                content.append(f"### 句子 {i}")
                content.append("")
                content.append(f"**英文**: {sentence['sentence']}")
                content.append(f"**中文**: {sentence['translation']}")
                if sentence.get('target_words'):
                    content.append(f"**目标单词**: {', '.join(sentence['target_words'])}")
                if sentence.get('morphology_points'):
                    content.append(f"**词法点**: {', '.join(sentence['morphology_points'])}")
                if sentence.get('syntax_structure'):
                    content.append(f"**句法结构**: {sentence['syntax_structure']}")
                content.append("")
    
    def _add_exercises_markdown(self, content, exercises_content):
        """添加练习题Markdown内容"""
        if exercises_content.get('practice_exercises'):
            # 按题型分组
            type_groups = {}
            for exercise in exercises_content['practice_exercises']:
                ex_type = exercise['type']
                if ex_type not in type_groups:
                    type_groups[ex_type] = []
                type_groups[ex_type].append(exercise)
            
            for ex_type, exercises in type_groups.items():
                type_name = {"choice": "选择题", "translation": "翻译题", "fill_blank": "填空题"}.get(ex_type, ex_type)
                content.append(f"### {type_name} ({len(exercises)}个)")
                content.append("")
                
                for i, exercise in enumerate(exercises, 1):
                    content.append(f"#### {i}. {exercise['question']}")
                    content.append("")
                    
                    if exercise['type'] == 'choice':
                        for j, option in enumerate(exercise['options']):
                            content.append(f"{chr(65+j)}. {option}")
                        content.append(f"**答案**: {exercise['correct_answer']}")
                    
                    elif exercise['type'] == 'translation':
                        content.append(f"**中文**: {exercise['chinese_text']}")
                        content.append(f"**英文**: {exercise['english_text']}")
                    
                    elif exercise['type'] == 'fill_blank':
                        content.append(f"**句子**: {exercise['sentence']}")
                        content.append(f"**答案**: {exercise['answer']}")
                    
                    content.append(f"**解析**: {exercise['explanation']}")
                    content.append("")
    
    def generate_and_display(self, target_date: str = None, plan_id: str = None):
        """生成并显示完整学习文档"""
        try:
            # 生成学习内容
            daily_content = self.generate_daily_learning_document(target_date, plan_id)
            
            # 生成Word文档
            word_path = self.generate_word_document(daily_content)
            
            print(f"✅ 完整学习文档已生成: {word_path}")
            print(f"📅 学习日期: {daily_content['date']}")
            print(f"📚 学习阶段: {daily_content['stage']}")
            print(f"📖 单词数量: {daily_content['words'].get('total_words', 0)}")
            print(f"📝 词法项目: {daily_content['morphology'].get('total_items', 0)}")
            print(f"📋 句法项目: {daily_content['syntax'].get('total_items', 0)}")
            print(f"💬 练习句子: {daily_content['sentences'].get('total_sentences', 0)}")
            print(f"✏️ 练习题: {daily_content['exercises'].get('total_exercises', 0)}")
            
        except Exception as e:
            print(f"❌ 生成学习文档失败: {e}")


def main():
    """主函数"""
    generator = DailyLearningDocumentGenerator()
    
    print("📚 每日学习内容文档生成器")
    print("=" * 50)
    
    # 获取学习计划列表
    plans = generator.plan_reader.list_plans()
    if not plans:
        print("❌ 没有找到学习计划，请先生成学习计划")
        return
    
    print("📋 可用的学习计划:")
    for i, plan in enumerate(plans, 1):
        print(f"{i}. {plan['name']} ({plan['date']})")
    
    # 选择学习计划
    while True:
        try:
            choice = input(f"\n请选择学习计划 (1-{len(plans)}): ").strip()
            plan_index = int(choice) - 1
            if 0 <= plan_index < len(plans):
                selected_plan = plans[plan_index]
                break
            else:
                print("❌ 无效选择，请重新输入")
        except ValueError:
            print("❌ 请输入有效数字")
    
    # 获取目标日期
    target_date = input("请输入学习日期 (格式: YYYY-MM-DD，直接回车使用今天): ").strip()
    if not target_date:
        target_date = datetime.now().strftime("%Y-%m-%d")
    
    # 生成文档
    generator.generate_and_display(target_date, selected_plan['id'])


if __name__ == "__main__":
    main()
