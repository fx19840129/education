#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自定义学习计划生成器
支持用户指定计划时长、每日学习时间，自动规划学习内容
"""

import sys
import os
import argparse
from datetime import datetime, timedelta
import json
import math

# 添加模块路径
sys.path.append(os.path.join(os.path.dirname(__file__), 'plan_modules'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'word_learning_modules'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'grammar_modules'))

try:
    from daily_content_generator import DailyContentGenerator
    from plan_document_generator import PlanDocumentGenerator
    from word_database import WordDatabase
    from grammar_config_loader import GrammarConfigLoader
except ImportError as e:
    print(f"❌ 模块导入失败: {e}")
    print("请确保在正确的目录下运行此脚本")
    sys.exit(1)

class CustomPlanGenerator:
    """自定义学习计划生成器"""
    
    def __init__(self):
        """初始化生成器"""
        print("📚 初始化自定义学习计划生成器...")
        
        self.daily_generator = DailyContentGenerator()
        self.doc_generator = PlanDocumentGenerator()
        self.word_db = WordDatabase()
        self.grammar_loader = GrammarConfigLoader()
        
        # 学习配置模板
        self.learning_templates = self._setup_learning_templates()
        
        # 学习阶段配置
        self.learning_stages = self._setup_learning_stages()
        
        print("✅ 初始化完成")
    
    def _setup_learning_templates(self):
        """设置学习配置模板"""
        return {
            "intensive": {
                "name": "强化学习",
                "daily_words": 15,
                "daily_grammar": 2,
                "daily_exercises": 30,
                "study_time_per_day": 60,  # 分钟
                "description": "高强度学习，适合有充足时间的学习者"
            },
            "standard": {
                "name": "标准学习",
                "daily_words": 10,
                "daily_grammar": 1,
                "daily_exercises": 20,
                "study_time_per_day": 30,  # 分钟
                "description": "平衡的学习强度，适合大多数学习者"
            },
            "light": {
                "name": "轻松学习",
                "daily_words": 5,
                "daily_grammar": 1,
                "daily_exercises": 10,
                "study_time_per_day": 15,  # 分钟
                "description": "轻松的学习节奏，适合时间有限的学习者"
            },
            "custom": {
                "name": "自定义学习",
                "daily_words": 0,  # 用户自定义
                "daily_grammar": 0,  # 用户自定义
                "daily_exercises": 0,  # 用户自定义
                "study_time_per_day": 0,  # 用户自定义
                "description": "完全自定义的学习计划"
            }
        }
    
    def _setup_learning_stages(self):
        """设置学习阶段配置"""
        return {
            "beginner": {
                "name": "初学者阶段",
                "description": "英语学习入门，重点掌握基础词汇和语法",
                "word_focus": "elementary",
                "grammar_focus": "elementary",
                "difficulty": "easy",
                "word_ratio": 0.8,  # 80%小学词汇
                "grammar_ratio": 0.7,  # 70%小学语法
                "learning_goals": [
                    "掌握基础词汇500-800个",
                    "学会基本语法结构",
                    "建立英语学习兴趣",
                    "培养学习习惯"
                ]
            },
            "intermediate": {
                "name": "中级阶段",
                "description": "英语能力提升，扩展词汇量和语法应用",
                "word_focus": "balanced",
                "grammar_focus": "balanced",
                "difficulty": "medium",
                "word_ratio": 0.5,  # 50%小学，50%初中
                "grammar_ratio": 0.5,  # 50%小学，50%初中
                "learning_goals": [
                    "掌握词汇1000-1500个",
                    "熟练运用各种语法结构",
                    "提高听说读写能力",
                    "增强语言应用能力"
                ]
            },
            "advanced": {
                "name": "高级阶段",
                "description": "英语能力深化，重点掌握复杂语法和高级词汇",
                "word_focus": "middle_school",
                "grammar_focus": "middle_school",
                "difficulty": "hard",
                "word_ratio": 0.2,  # 20%小学，80%初中
                "grammar_ratio": 0.3,  # 30%小学，70%初中
                "learning_goals": [
                    "掌握词汇1500-2000个",
                    "精通复杂语法结构",
                    "提高语言表达准确性",
                    "培养英语思维"
                ]
            },
            "comprehensive": {
                "name": "综合阶段",
                "description": "全面复习巩固，综合运用所有知识点",
                "word_focus": "mixed",
                "grammar_focus": "mixed",
                "difficulty": "mixed",
                "word_ratio": 0.4,  # 40%小学，60%初中
                "grammar_ratio": 0.4,  # 40%小学，60%初中
                "learning_goals": [
                    "全面复习已学内容",
                    "查漏补缺薄弱环节",
                    "提高综合应用能力",
                    "准备进阶学习"
                ]
            }
        }
    
    def calculate_learning_plan(self, total_days, daily_minutes, custom_config=None, stage="intermediate"):
        """计算学习计划"""
        print(f"📊 计算学习计划...")
        print(f"   计划时长: {total_days} 天")
        print(f"   每日学习时间: {daily_minutes} 分钟")
        
        # 获取学习阶段配置
        stage_config = self.learning_stages.get(stage, self.learning_stages["intermediate"])
        print(f"   学习阶段: {stage_config['name']}")
        
        # 根据每日学习时间选择学习模板
        if custom_config:
            template = custom_config
            template_name = "自定义"
        elif daily_minutes >= 50:
            template = self.learning_templates["intensive"]
            template_name = "强化学习"
        elif daily_minutes >= 25:
            template = self.learning_templates["standard"]
            template_name = "标准学习"
        else:
            template = self.learning_templates["light"]
            template_name = "轻松学习"
        
        print(f"   学习模式: {template_name}")
        
        # 计算学习内容
        total_words = template["daily_words"] * total_days
        total_grammar = template["daily_grammar"] * total_days
        total_exercises = template["daily_exercises"] * total_days
        total_study_time = daily_minutes * total_days
        
        # 根据学习阶段调整词汇和语法分布
        word_distribution = self._calculate_stage_word_distribution(total_words, stage_config)
        grammar_distribution = self._calculate_stage_grammar_distribution(total_grammar, stage_config)
        
        # 生成计划数据
        plan_data = {
            "metadata": {
                "plan_name": f"{stage_config['name']}{total_days}天学习计划",
                "total_days": total_days,
                "daily_minutes": daily_minutes,
                "template_name": template_name,
                "stage_name": stage_config["name"],
                "generation_time": datetime.now().isoformat()
            },
            "stage_info": {
                "name": stage_config["name"],
                "description": stage_config["description"],
                "difficulty": stage_config["difficulty"],
                "learning_goals": stage_config["learning_goals"]
            },
            "learning_config": {
                "daily_words": template["daily_words"],
                "daily_grammar": template["daily_grammar"],
                "daily_exercises": template["daily_exercises"],
                "study_time_per_day": daily_minutes,
                "description": template["description"]
            },
            "statistics": {
                "total_words": total_words,
                "total_grammar_points": total_grammar,
                "total_exercises": total_exercises,
                "total_study_time_minutes": total_study_time,
                "total_study_time_hours": round(total_study_time / 60, 1),
                "average_words_per_day": template["daily_words"],
                "average_grammar_per_day": template["daily_grammar"],
                "average_exercises_per_day": template["daily_exercises"]
            },
            "distribution": {
                "grammar": grammar_distribution,
                "words": word_distribution
            },
            "weekly_breakdown": self._calculate_weekly_breakdown(total_days, template)
        }
        
        return plan_data
    
    def _calculate_stage_word_distribution(self, total_words, stage_config):
        """根据学习阶段计算词汇分布"""
        word_ratio = stage_config["word_ratio"]
        
        if stage_config["word_focus"] == "elementary":
            elementary_words = total_words
            middle_school_words = 0
        elif stage_config["word_focus"] == "middle_school":
            elementary_words = 0
            middle_school_words = total_words
        elif stage_config["word_focus"] == "balanced":
            elementary_words = int(total_words * 0.5)
            middle_school_words = total_words - elementary_words
        else:  # mixed
            elementary_words = int(total_words * word_ratio)
            middle_school_words = total_words - elementary_words
        
        return {
            "elementary": {
                "count": elementary_words,
                "percentage": round(elementary_words / total_words * 100, 1) if total_words > 0 else 0
            },
            "middle_school": {
                "count": middle_school_words,
                "percentage": round(middle_school_words / total_words * 100, 1) if total_words > 0 else 0
            }
        }
    
    def _calculate_stage_grammar_distribution(self, total_grammar, stage_config):
        """根据学习阶段计算语法分布"""
        grammar_ratio = stage_config["grammar_ratio"]
        
        # 基础语法点列表
        elementary_grammar = [
            "be动词用法-基础用法", "一般现在时-基础用法", "名词单复数-基础规则",
            "形容词-基础用法", "there_be-基础用法", "现在进行时-基础用法",
            "过去时-基础用法", "形容词比较级-基础规则", "情态动词-基础用法"
        ]
        
        middle_school_grammar = [
            "现在完成时-基础用法", "被动语态-基础用法", "定语从句-基础用法",
            "条件句-基础用法", "过去进行时-基础用法", "间接引语-基础用法",
            "现在完成时-持续用法", "被动语态-时态变化", "定语从句-关系代词",
            "条件句-虚拟语气", "非谓语动词-基础用法", "倒装句-基础用法",
            "强调句-基础用法", "主谓一致-复杂情况", "虚拟语气-完整体系"
        ]
        
        # 实际可用的语法点数量
        max_elementary = len(elementary_grammar)  # 9个
        max_middle_school = len(middle_school_grammar)  # 15个
        max_total = max_elementary + max_middle_school  # 24个
        
        if stage_config["grammar_focus"] == "elementary":
            elementary_count = min(total_grammar, max_elementary)
            middle_school_count = 0
        elif stage_config["grammar_focus"] == "middle_school":
            elementary_count = 0
            middle_school_count = min(total_grammar, max_middle_school)
        elif stage_config["grammar_focus"] == "balanced":
            elementary_count = min(max_elementary, total_grammar // 2)
            middle_school_count = min(max_middle_school, total_grammar - elementary_count)
        else:  # mixed
            elementary_count = min(max_elementary, int(total_grammar * grammar_ratio))
            middle_school_count = min(max_middle_school, total_grammar - elementary_count)
        
        # 如果总语法点数超过实际可用数量，需要重复学习
        if total_grammar > max_total:
            remaining = total_grammar - max_total
            if remaining > 0:
                # 按阶段比例重复分配
                elementary_repeat = int(remaining * grammar_ratio)
                middle_school_repeat = remaining - elementary_repeat
                elementary_count += elementary_repeat
                middle_school_count += middle_school_repeat
        
        return {
            "elementary": {
                "count": elementary_count,
                "topics": elementary_grammar[:min(elementary_count, max_elementary)] if elementary_count > 0 else []
            },
            "middle_school": {
                "count": middle_school_count,
                "topics": middle_school_grammar[:min(middle_school_count, max_middle_school)] if middle_school_count > 0 else []
            }
        }
    
    def _calculate_grammar_distribution(self, total_grammar):
        """计算语法点分布"""
        # 基础语法点列表
        elementary_grammar = [
            "be动词用法-基础用法", "一般现在时-基础用法", "名词单复数-基础规则",
            "形容词-基础用法", "there_be-基础用法", "现在进行时-基础用法",
            "过去时-基础用法", "形容词比较级-基础规则", "情态动词-基础用法"
        ]
        
        middle_school_grammar = [
            "现在完成时-基础用法", "被动语态-基础用法", "定语从句-基础用法",
            "条件句-基础用法", "过去进行时-基础用法", "间接引语-基础用法",
            "现在完成时-持续用法", "被动语态-时态变化", "定语从句-关系代词",
            "条件句-虚拟语气", "非谓语动词-基础用法", "倒装句-基础用法",
            "强调句-基础用法", "主谓一致-复杂情况", "虚拟语气-完整体系"
        ]
        
        # 实际可用的语法点数量
        max_elementary = len(elementary_grammar)  # 9个
        max_middle_school = len(middle_school_grammar)  # 15个
        max_total = max_elementary + max_middle_school  # 24个
        
        # 如果总语法点数超过实际可用数量，需要重复学习
        if total_grammar <= max_total:
            # 在可用范围内分配
            if total_grammar <= max_elementary:
                # 主要学习小学语法
                elementary_count = total_grammar
                middle_school_count = 0
            else:
                # 平衡分配
                elementary_count = min(max_elementary, total_grammar // 2)
                middle_school_count = min(max_middle_school, total_grammar - elementary_count)
        else:
            # 超过实际可用数量，需要重复学习
            # 先分配完所有可用语法点
            elementary_count = max_elementary
            middle_school_count = max_middle_school
            remaining = total_grammar - max_total
            
            # 剩余语法点按比例重复分配
            if remaining > 0:
                # 按7:3的比例重复分配（小学:初中）
                elementary_repeat = int(remaining * 0.7)
                middle_school_repeat = remaining - elementary_repeat
                elementary_count += elementary_repeat
                middle_school_count += middle_school_repeat
        
        return {
            "elementary": {
                "count": elementary_count,
                "topics": elementary_grammar[:min(elementary_count, max_elementary)] if elementary_count > 0 else []
            },
            "middle_school": {
                "count": middle_school_count,
                "topics": middle_school_grammar[:min(middle_school_count, max_middle_school)] if middle_school_count > 0 else []
            }
        }
    
    def _calculate_word_distribution(self, total_words):
        """计算词汇分布"""
        # 根据总词汇数分配小学和初中词汇
        if total_words <= 100:
            elementary_words = total_words
            middle_school_words = 0
        elif total_words <= 300:
            elementary_words = min(200, total_words * 0.7)
            middle_school_words = total_words - elementary_words
        else:
            elementary_words = min(300, total_words * 0.4)
            middle_school_words = total_words - elementary_words
        
        return {
            "elementary": {
                "count": int(elementary_words),
                "percentage": round(elementary_words / total_words * 100, 1) if total_words > 0 else 0
            },
            "middle_school": {
                "count": int(middle_school_words),
                "percentage": round(middle_school_words / total_words * 100, 1) if total_words > 0 else 0
            }
        }
    
    def _calculate_weekly_breakdown(self, total_days, template):
        """计算每周学习安排"""
        weeks = []
        remaining_days = total_days
        
        week_number = 1
        while remaining_days > 0:
            week_days = min(7, remaining_days)
            
            # 周末减少学习量
            weekend_days = min(2, week_days - 5) if week_days > 5 else 0
            weekday_days = week_days - weekend_days
            
            week_words = (weekday_days * template["daily_words"] + 
                         weekend_days * max(1, template["daily_words"] // 2))
            week_grammar = (weekday_days * template["daily_grammar"] + 
                           weekend_days * max(1, template["daily_grammar"] // 2))
            week_exercises = (weekday_days * template["daily_exercises"] + 
                             weekend_days * max(5, template["daily_exercises"] // 2))
            
            weeks.append({
                "week": week_number,
                "days": week_days,
                "words": week_words,
                "grammar": week_grammar,
                "exercises": week_exercises
            })
            
            remaining_days -= week_days
            week_number += 1
        
        return weeks
    
    def generate_custom_plan(self, total_days, daily_minutes, start_date=None, 
                           custom_words=None, custom_grammar=None, custom_exercises=None,
                           stage="intermediate", output_dir="custom_plans"):
        """生成自定义学习计划"""
        if start_date is None:
            start_date = datetime.now()
        elif isinstance(start_date, str):
            start_date = datetime.strptime(start_date, "%Y-%m-%d")
        
        # 自定义配置
        custom_config = None
        if custom_words or custom_grammar or custom_exercises:
            custom_config = {
                "name": "自定义学习",
                "daily_words": custom_words or 10,
                "daily_grammar": custom_grammar or 1,
                "daily_exercises": custom_exercises or 20,
                "study_time_per_day": daily_minutes,
                "description": "完全自定义的学习计划"
            }
        
        # 计算学习计划
        plan_data = self.calculate_learning_plan(total_days, daily_minutes, custom_config, stage)
        
        # 添加时间信息
        plan_data["metadata"]["start_date"] = start_date.strftime("%Y-%m-%d")
        plan_data["metadata"]["end_date"] = (start_date + timedelta(days=total_days-1)).strftime("%Y-%m-%d")
        
        # 保存文件
        files = self._save_custom_plan(plan_data, output_dir)
        
        return {
            "plan_data": plan_data,
            "files": files
        }
    
    def list_stages(self):
        """列出所有可用学习阶段"""
        print("📋 可用学习阶段:")
        print("-" * 50)
        for stage_id, stage_config in self.learning_stages.items():
            print(f"• {stage_id}: {stage_config['name']}")
            print(f"  描述: {stage_config['description']}")
            print(f"  难度: {stage_config['difficulty']}")
            print(f"  词汇重点: {stage_config['word_focus']}")
            print(f"  语法重点: {stage_config['grammar_focus']}")
            print(f"  学习目标:")
            for goal in stage_config['learning_goals']:
                print(f"    - {goal}")
            print()
    
    def _save_custom_plan(self, plan_data, output_dir):
        """保存自定义计划到文件"""
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        plan_name = f"custom_plan_{plan_data['metadata']['total_days']}days_{timestamp}"
        
        # 保存JSON格式
        json_file = f"{output_dir}/{plan_name}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(plan_data, f, ensure_ascii=False, indent=2)
        
        # 生成Word文档
        word_file = f"{output_dir}/{plan_name}.docx"
        self._generate_custom_plan_document(plan_data, word_file)
        
        # 生成文本总结
        txt_file = f"{output_dir}/{plan_name}_summary.txt"
        self._generate_custom_plan_summary(plan_data, txt_file)
        
        return {
            "json_file": json_file,
            "word_file": word_file,
            "txt_file": txt_file
        }
    
    def _generate_custom_plan_document(self, plan_data, output_file):
        """生成自定义计划的Word文档"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(f'{plan_data["metadata"]["plan_name"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        metadata = plan_data["metadata"]
        stats = plan_data["statistics"]
        
        info_data = [
            ("计划时长", f"{metadata['total_days']} 天"),
            ("开始日期", metadata["start_date"]),
            ("结束日期", metadata["end_date"]),
            ("每日学习时间", f"{metadata['daily_minutes']} 分钟"),
            ("学习模式", metadata["template_name"]),
            ("总学习时间", f"{stats['total_study_time_hours']} 小时")
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # 学习内容统计
        doc.add_heading('学习内容统计', level=1)
        
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ("总词汇量", f"{stats['total_words']} 个"),
            ("总语法点", f"{stats['total_grammar_points']} 个"),
            ("总练习题", f"{stats['total_exercises']} 题"),
            ("日均学习量", f"{stats['average_words_per_day']}词 + {stats['average_grammar_per_day']}语法 + {stats['average_exercises_per_day']}题")
        ]
        
        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
        
        # 词汇分布
        doc.add_heading('词汇分布', level=1)
        word_dist = plan_data["distribution"]["words"]
        
        p = doc.add_paragraph()
        p.add_run("小学词汇：").bold = True
        p.add_run(f" {word_dist['elementary']['count']} 个 ({word_dist['elementary']['percentage']}%)")
        
        p = doc.add_paragraph()
        p.add_run("初中词汇：").bold = True
        p.add_run(f" {word_dist['middle_school']['count']} 个 ({word_dist['middle_school']['percentage']}%)")
        
        # 语法分布
        doc.add_heading('语法分布', level=1)
        grammar_dist = plan_data["distribution"]["grammar"]
        
        p = doc.add_paragraph()
        p.add_run("小学语法：").bold = True
        p.add_run(f" {grammar_dist['elementary']['count']} 个")
        if grammar_dist['elementary']['topics']:
            for topic in grammar_dist['elementary']['topics']:
                doc.add_paragraph(f"  • {topic}", style='List Bullet')
        
        p = doc.add_paragraph()
        p.add_run("初中语法：").bold = True
        p.add_run(f" {grammar_dist['middle_school']['count']} 个")
        if grammar_dist['middle_school']['topics']:
            for topic in grammar_dist['middle_school']['topics']:
                doc.add_paragraph(f"  • {topic}", style='List Bullet')
        
        # 每周安排
        doc.add_heading('每周学习安排', level=1)
        
        for week in plan_data["weekly_breakdown"]:
            doc.add_heading(f'第{week["week"]}周', level=2)
            
            p = doc.add_paragraph()
            p.add_run("学习天数：").bold = True
            p.add_run(f" {week['days']} 天")
            p.add_run(" | 词汇：").bold = True
            p.add_run(f" {week['words']} 个")
            p.add_run(" | 语法：").bold = True
            p.add_run(f" {week['grammar']} 个")
            p.add_run(" | 练习：").bold = True
            p.add_run(f" {week['exercises']} 题")
        
        # 学习建议
        doc.add_heading('学习建议', level=1)
        suggestions = [
            f"本计划采用{metadata['template_name']}模式，适合{plan_data['learning_config']['description']}",
            f"每天坚持学习{metadata['daily_minutes']}分钟，保持学习连续性",
            f"建议每天学习{stats['average_words_per_day']}个单词，{stats['average_grammar_per_day']}个语法点",
            f"完成{stats['average_exercises_per_day']}道练习题巩固所学内容",
            "定期复习前面学过的内容，加强记忆效果",
            "可以根据个人进度适当调整学习量",
            f"完成本计划后，您将掌握{stats['total_words']}个单词和{stats['total_grammar_points']}个语法点"
        ]
        
        for suggestion in suggestions:
            p = doc.add_paragraph(suggestion, style='List Bullet')
        
        doc.save(output_file)
    
    def _generate_custom_plan_summary(self, plan_data, output_file):
        """生成自定义计划文本总结"""
        with open(output_file, 'w', encoding='utf-8') as f:
            metadata = plan_data["metadata"]
            stats = plan_data["statistics"]
            word_dist = plan_data["distribution"]["words"]
            grammar_dist = plan_data["distribution"]["grammar"]
            
            f.write(f"{metadata['plan_name']}总结\n")
            f.write("=" * 50 + "\n\n")
            
            f.write("基本信息:\n")
            f.write("-" * 20 + "\n")
            f.write(f"计划时长: {metadata['total_days']} 天\n")
            f.write(f"开始日期: {metadata['start_date']}\n")
            f.write(f"结束日期: {metadata['end_date']}\n")
            f.write(f"每日学习时间: {metadata['daily_minutes']} 分钟\n")
            f.write(f"学习模式: {metadata['template_name']}\n")
            f.write(f"总学习时间: {stats['total_study_time_hours']} 小时\n\n")
            
            f.write("学习内容:\n")
            f.write("-" * 20 + "\n")
            f.write(f"总词汇量: {stats['total_words']} 个\n")
            f.write(f"  小学词汇: {word_dist['elementary']['count']} 个 ({word_dist['elementary']['percentage']}%)\n")
            f.write(f"  初中词汇: {word_dist['middle_school']['count']} 个 ({word_dist['middle_school']['percentage']}%)\n")
            f.write(f"总语法点: {stats['total_grammar_points']} 个\n")
            f.write(f"  小学语法: {grammar_dist['elementary']['count']} 个\n")
            f.write(f"  初中语法: {grammar_dist['middle_school']['count']} 个\n")
            f.write(f"总练习题: {stats['total_exercises']} 题\n\n")
            
            f.write("每日学习量:\n")
            f.write("-" * 20 + "\n")
            f.write(f"单词: {stats['average_words_per_day']} 个\n")
            f.write(f"语法: {stats['average_grammar_per_day']} 个\n")
            f.write(f"练习: {stats['average_exercises_per_day']} 题\n\n")
            
            f.write("每周安排:\n")
            f.write("-" * 20 + "\n")
            for week in plan_data["weekly_breakdown"]:
                f.write(f"第{week['week']}周: {week['days']}天, {week['words']}词, {week['grammar']}语法, {week['exercises']}题\n")
            
            f.write(f"\n学习成果:\n")
            f.write("-" * 20 + "\n")
            f.write(f"完成本计划后，您将掌握:\n")
            f.write(f"• {stats['total_words']} 个英语单词\n")
            f.write(f"• {stats['total_grammar_points']} 个语法知识点\n")
            f.write(f"• 完成 {stats['total_exercises']} 道练习题\n")
            f.write(f"• 累计学习 {stats['total_study_time_hours']} 小时\n")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description="自定义英语学习计划生成器")
    parser.add_argument("--days", "-d", type=int, help="计划时长（天数）")
    parser.add_argument("--minutes", "-m", type=int, help="每日学习时间（分钟）")
    parser.add_argument("--start-date", "-s", help="开始日期 (YYYY-MM-DD)")
    parser.add_argument("--words", type=int, help="自定义每日单词数")
    parser.add_argument("--grammar", type=int, help="自定义每日语法点数")
    parser.add_argument("--exercises", type=int, help="自定义每日练习题数")
    parser.add_argument("--stage", choices=['beginner', 'intermediate', 'advanced', 'comprehensive'], 
                       default='intermediate', help="学习阶段")
    parser.add_argument("--output-dir", "-o", default="custom_plans", help="输出目录")
    parser.add_argument("--list-stages", action="store_true", help="列出所有可用学习阶段")
    
    args = parser.parse_args()
    
    print("📚 === 自定义英语学习计划生成器 ===")
    print("🎯 根据您的需求制定个性化学习计划")
    print()
    
    try:
        generator = CustomPlanGenerator()
        
        if args.list_stages:
            # 列出所有学习阶段
            generator.list_stages()
            
        elif args.days and args.minutes:
            # 生成自定义计划
            result = generator.generate_custom_plan(
                total_days=args.days,
                daily_minutes=args.minutes,
                start_date=args.start_date,
                custom_words=args.words,
                custom_grammar=args.grammar,
                custom_exercises=args.exercises,
                stage=args.stage,
                output_dir=args.output_dir
            )
            
            plan_data = result["plan_data"]
            files = result["files"]
            stats = plan_data["statistics"]
            stage_info = plan_data["stage_info"]
            
            print(f"\n🎉 {plan_data['metadata']['plan_name']}生成完成！")
            print(f"📄 详细计划: {files['word_file']}")
            print(f"📊 数据文件: {files['json_file']}")
            print(f"📋 总结文件: {files['txt_file']}")
            
            print(f"\n📈 学习计划概览:")
            print(f"   计划时长: {plan_data['metadata']['total_days']} 天")
            print(f"   每日学习时间: {plan_data['metadata']['daily_minutes']} 分钟")
            print(f"   学习模式: {plan_data['metadata']['template_name']}")
            print(f"   学习阶段: {stage_info['name']}")
            print(f"   总学习时间: {stats['total_study_time_hours']} 小时")
            
            print(f"\n📚 学习内容:")
            print(f"   总词汇量: {stats['total_words']} 个")
            print(f"   总语法点: {stats['total_grammar_points']} 个")
            print(f"   总练习题: {stats['total_exercises']} 题")
            
            word_dist = plan_data["distribution"]["words"]
            print(f"\n📖 词汇分布:")
            print(f"   小学词汇: {word_dist['elementary']['count']} 个 ({word_dist['elementary']['percentage']}%)")
            print(f"   初中词汇: {word_dist['middle_school']['count']} 个 ({word_dist['middle_school']['percentage']}%)")
            
            grammar_dist = plan_data["distribution"]["grammar"]
            print(f"\n📝 语法分布:")
            print(f"   小学语法: {grammar_dist['elementary']['count']} 个")
            print(f"   初中语法: {grammar_dist['middle_school']['count']} 个")
            
            print(f"\n🎯 学习目标:")
            for goal in stage_info['learning_goals']:
                print(f"   • {goal}")
            
            print(f"\n🎯 学习成果:")
            print(f"   完成本计划后，您将掌握 {stats['total_words']} 个单词和 {stats['total_grammar_points']} 个语法点！")
            
        else:
            # 默认列出所有阶段
            print("💡 请指定要生成的计划参数，或使用 --list-stages 查看所有可用阶段")
            print("\n使用方法:")
            print("  python custom_plan_generator.py --list-stages  # 列出所有阶段")
            print("  python custom_plan_generator.py --days 30 --minutes 30  # 生成30天计划")
            print("  python custom_plan_generator.py --days 60 --minutes 20 --stage beginner  # 初学者阶段")
            print("  python custom_plan_generator.py --days 90 --minutes 45 --stage advanced  # 高级阶段")
    
    except Exception as e:
        print(f"❌ 生成失败: {e}")
        print("💡 请检查参数是否正确")


if __name__ == "__main__":
    main()
