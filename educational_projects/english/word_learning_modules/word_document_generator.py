#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
单词学习Word文档生成器
生成专业的Word文档用于单词学习
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from typing import List, Dict, Any
from exercise_generator import Exercise
from learning_plan_generator import LearningPlan
from word_database import WordInfo
import os
from datetime import datetime


class WordDocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.output_dir = "word_learning_details"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generate_word_exercises_document(self, words: List[WordInfo], exercises: List[Exercise], 
                                       level: str, difficulty: str = None) -> str:
        """生成单词练习Word文档"""
        doc = Document()
        self._setup_document_styles(doc)
        
        # 添加标题
        title = f"{'小学' if level == 'elementary' else '初中'}英语单词练习"
        if difficulty:
            difficulty_name = {"easy": "基础", "medium": "进阶", "hard": "高级"}.get(difficulty, difficulty)
            title += f"（{difficulty_name}）"
        
        doc.add_heading(title, 0)
        
        # 添加基本信息
        self._add_basic_info(doc, words, level, difficulty)
        
        # 添加单词列表
        self._add_word_list(doc, words)
        
        # 添加练习题
        self._add_exercises(doc, exercises)
        
        # 添加学习建议
        self._add_learning_tips(doc, level)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{level}_word_exercises_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_word_answers_document(self, exercises: List[Exercise], level: str) -> str:
        """生成单词练习答案Word文档"""
        doc = Document()
        self._setup_document_styles(doc)
        
        # 添加标题
        title = f"{'小学' if level == 'elementary' else '初中'}英语单词练习答案"
        doc.add_heading(title, 0)
        
        # 添加答案
        self._add_answers(doc, exercises)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{level}_word_answers_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_learning_plan_document(self, plan: LearningPlan) -> str:
        """生成学习计划Word文档"""
        doc = Document()
        self._setup_document_styles(doc)
        
        # 添加标题
        doc.add_heading(plan.title, 0)
        
        # 添加计划描述
        self._add_plan_description(doc, plan)
        
        # 添加学习目标
        self._add_learning_objectives(doc, plan)
        
        # 添加学习计划表
        self._add_study_schedule(doc, plan)
        
        # 添加评估标准
        self._add_assessment_criteria(doc, plan)
        
        # 添加学习建议
        self._add_learning_tips(doc, plan.level)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{plan.level}_learning_plan_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.2
        
        # 设置标题样式
        heading1 = doc.styles['Heading 1']
        heading1.font.name = '宋体'
        heading1.font.size = Pt(14)
        heading1.paragraph_format.line_spacing = 1.2
        
        heading2 = doc.styles['Heading 2']
        heading2.font.name = '宋体'
        heading2.font.size = Pt(12)
        heading2.paragraph_format.line_spacing = 1.2
        
        heading3 = doc.styles['Heading 3']
        heading3.font.name = '宋体'
        heading3.font.size = Pt(10.5)
        heading3.paragraph_format.line_spacing = 1.2
    
    def _add_basic_info(self, doc: Document, words: List[WordInfo], level: str, difficulty: str):
        """添加基本信息"""
        doc.add_heading("基本信息", 1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.style = 'Table Grid'
        
        # 设置表格内容
        info_data = [
            ("年级级别", "小学" if level == "elementary" else "初中"),
            ("难度级别", {"easy": "简单", "medium": "中等", "hard": "困难"}.get(difficulty, "综合")),
            ("单词数量", str(len(words))),
            ("生成时间", datetime.now().strftime("%Y年%m月%d日"))
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
            
            # 设置单元格样式
            for j in range(2):
                cell = info_table.cell(i, j)
                cell.paragraphs[0].runs[0].font.name = '宋体'
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.2
        
        doc.add_paragraph()
    
    def _add_word_list(self, doc: Document, words: List[WordInfo]):
        """添加单词列表"""
        doc.add_heading("单词列表", 1)
        
        # 创建单词表格
        word_table = doc.add_table(rows=1, cols=5)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        word_table.style = 'Table Grid'
        
        # 设置表头
        header_cells = word_table.rows[0].cells
        headers = ["单词", "发音", "词性", "中文意思", "例句"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.name = '宋体'
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
            header_cells[i].paragraphs[0].paragraph_format.line_spacing = 1.2
        
        # 添加单词数据
        for word in words:
            row_cells = word_table.add_row().cells
            row_cells[0].text = word.word
            row_cells[1].text = word.pronunciation
            row_cells[2].text = word.part_of_speech
            row_cells[3].text = word.chinese_meaning
            row_cells[4].text = word.example_sentence
            
            # 设置单元格样式
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.name = '宋体'
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.2
        
        doc.add_paragraph()
    
    def _add_exercises(self, doc: Document, exercises: List[Exercise]):
        """添加练习题"""
        doc.add_heading("练习题", 1)
        
        for i, exercise in enumerate(exercises, 1):
            # 添加题目
            question_para = doc.add_paragraph()
            question_para.add_run(f"第{i}题：").bold = True
            question_para.add_run(f" {exercise.question}")
            question_para.paragraph_format.line_spacing = 1.2
            question_para.paragraph_format.space_after = Pt(3)
            
            # 添加选项（如果有）
            if exercise.options:
                options_para = doc.add_paragraph()
                options_para.add_run("选项：").bold = True
                options_para.paragraph_format.line_spacing = 1.2
                options_para.paragraph_format.space_after = Pt(3)
                
                for j, option in enumerate(exercise.options):
                    option_para = doc.add_paragraph(f"{chr(65+j)}. {option}", style='List Bullet')
                    option_para.paragraph_format.line_spacing = 1.2
                    option_para.paragraph_format.space_after = Pt(1)
            
            # 添加答题空间
            answer_para = doc.add_paragraph()
            answer_para.add_run("答案：")
            answer_para.paragraph_format.line_spacing = 1.2
            answer_para.paragraph_format.space_after = Pt(6)
            
            # 如果不是最后一道题，添加分隔线
            if i < len(exercises):
                doc.add_paragraph("_" * 50)
                doc.add_paragraph()
    
    def _add_answers(self, doc: Document, exercises: List[Exercise]):
        """添加答案"""
        doc.add_heading("练习答案", 1)
        
        for i, exercise in enumerate(exercises, 1):
            # 添加题目
            question_para = doc.add_paragraph()
            question_para.add_run(f"第{i}题：").bold = True
            question_para.add_run(f" {exercise.question}")
            question_para.paragraph_format.line_spacing = 1.2
            question_para.paragraph_format.space_after = Pt(3)
            
            # 添加选项（如果有）
            if exercise.options:
                options_para = doc.add_paragraph()
                options_para.add_run("选项：").bold = True
                options_para.paragraph_format.line_spacing = 1.2
                options_para.paragraph_format.space_after = Pt(3)
                
                for j, option in enumerate(exercise.options):
                    option_para = doc.add_paragraph(f"{chr(65+j)}. {option}", style='List Bullet')
                    option_para.paragraph_format.line_spacing = 1.2
                    option_para.paragraph_format.space_after = Pt(1)
            
            # 添加正确答案
            answer_para = doc.add_paragraph()
            answer_para.add_run("正确答案：").bold = True
            answer_para.add_run(f" {exercise.correct_answer}")
            answer_para.paragraph_format.line_spacing = 1.2
            answer_para.paragraph_format.space_after = Pt(3)
            
            # 添加解析
            if exercise.explanation:
                explanation_para = doc.add_paragraph()
                explanation_para.add_run("解析：").bold = True
                explanation_para.add_run(f" {exercise.explanation}")
                explanation_para.paragraph_format.line_spacing = 1.2
                explanation_para.paragraph_format.space_after = Pt(6)
            
            # 如果不是最后一道题，添加分隔线
            if i < len(exercises):
                doc.add_paragraph("_" * 50)
                doc.add_paragraph()
    
    def _add_plan_description(self, doc: Document, plan: LearningPlan):
        """添加计划描述"""
        doc.add_heading("计划描述", 1)
        
        desc_para = doc.add_paragraph(plan.description)
        desc_para.paragraph_format.line_spacing = 1.2
        desc_para.paragraph_format.space_after = Pt(6)
        
        # 添加计划统计信息
        stats_para = doc.add_paragraph()
        stats_para.add_run("计划统计：").bold = True
        stats_para.paragraph_format.line_spacing = 1.2
        stats_para.paragraph_format.space_after = Pt(3)
        
        stats_list = [
            f"• 学习时长：{plan.duration}天",
            f"• 每日单词数：{plan.daily_words}个",
            f"• 总单词数：{plan.total_words}个",
            f"• 难度级别：{plan.difficulty}",
            f"• 学习分类：{', '.join(plan.categories) if plan.categories else '综合'}"
        ]
        
        for stat in stats_list:
            stat_para = doc.add_paragraph(stat)
            stat_para.paragraph_format.line_spacing = 1.2
            stat_para.paragraph_format.space_after = Pt(1)
    
    def _add_learning_objectives(self, doc: Document, plan: LearningPlan):
        """添加学习目标"""
        doc.add_heading("学习目标", 1)
        
        for objective in plan.learning_objectives:
            obj_para = doc.add_paragraph(f"• {objective}")
            obj_para.paragraph_format.line_spacing = 1.2
            obj_para.paragraph_format.space_after = Pt(1)
    
    def _add_study_schedule(self, doc: Document, plan: LearningPlan):
        """添加学习计划表"""
        doc.add_heading("学习计划表", 1)
        
        for day_plan in plan.study_schedule:
            # 添加日期标题
            day_heading = doc.add_heading(f"第{day_plan['day']}天", 2)
            day_heading.paragraph_format.line_spacing = 1.2
            day_heading.paragraph_format.space_after = Pt(6)
            
            # 添加学习方法
            method_para = doc.add_paragraph()
            method_para.add_run("学习方法：").bold = True
            method_para.add_run(f" {day_plan['learning_method']}")
            method_para.paragraph_format.line_spacing = 1.2
            method_para.paragraph_format.space_after = Pt(3)
            
            # 添加重点技能
            if day_plan['focus_skills']:
                skills_para = doc.add_paragraph()
                skills_para.add_run("重点技能：").bold = True
                skills_para.add_run(f" {', '.join(day_plan['focus_skills'])}")
                skills_para.paragraph_format.line_spacing = 1.2
                skills_para.paragraph_format.space_after = Pt(3)
            
            # 添加练习活动
            if day_plan['practice_activities']:
                activities_para = doc.add_paragraph()
                activities_para.add_run("练习活动：").bold = True
                activities_para.paragraph_format.line_spacing = 1.2
                activities_para.paragraph_format.space_after = Pt(3)
                
                for activity in day_plan['practice_activities']:
                    activity_para = doc.add_paragraph(f"• {activity}")
                    activity_para.paragraph_format.line_spacing = 1.2
                    activity_para.paragraph_format.space_after = Pt(1)
            
            # 添加今日单词
            words_para = doc.add_paragraph()
            words_para.add_run("今日单词：").bold = True
            words_para.paragraph_format.line_spacing = 1.2
            words_para.paragraph_format.space_after = Pt(3)
            
            for word_info in day_plan['words']:
                word_para = doc.add_paragraph(f"• {word_info['word']} - {word_info['meaning']} ({word_info['pronunciation']})")
                word_para.paragraph_format.line_spacing = 1.2
                word_para.paragraph_format.space_after = Pt(1)
            
            # 添加复习单词
            if day_plan['review_words']:
                review_para = doc.add_paragraph()
                review_para.add_run("复习单词：").bold = True
                review_para.add_run(f" {', '.join(day_plan['review_words'])}")
                review_para.paragraph_format.line_spacing = 1.2
                review_para.paragraph_format.space_after = Pt(6)
            
            # 添加分隔线
            if day_plan['day'] < len(plan.study_schedule):
                doc.add_paragraph("_" * 50)
                doc.add_paragraph()
    
    def _add_assessment_criteria(self, doc: Document, plan: LearningPlan):
        """添加评估标准"""
        doc.add_heading("评估标准", 1)
        
        for level, criteria in plan.assessment_criteria.items():
            level_para = doc.add_paragraph()
            level_para.add_run(f"{level}：").bold = True
            level_para.add_run(f" {criteria}")
            level_para.paragraph_format.line_spacing = 1.2
            level_para.paragraph_format.space_after = Pt(3)
    
    def _add_learning_tips(self, doc: Document, level: str):
        """添加学习建议"""
        doc.add_heading("学习建议", 1)
        
        tips = [
            "每天坚持学习，保持学习的连续性",
            "多听、多说、多读、多写，全方位练习",
            "制作单词卡片，随时复习",
            "将单词放入句子中学习，加深理解",
            "定期复习已学单词，防止遗忘",
            "寻找学习伙伴，互相督促和练习"
        ]
        
        if level == "elementary":
            tips.extend([
                "通过游戏和歌曲学习，增加趣味性",
                "家长可以参与学习过程，给予鼓励"
            ])
        else:
            tips.extend([
                "学习词根词缀，提高词汇量",
                "阅读英文文章，在语境中学习单词",
                "使用单词学习APP，提高学习效率"
            ])
        
        for tip in tips:
            tip_para = doc.add_paragraph(f"• {tip}")
            tip_para.paragraph_format.line_spacing = 1.2
            tip_para.paragraph_format.space_after = Pt(1)
