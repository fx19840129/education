#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学习计划文档生成器
生成包含具体单词和句子的学习计划文档
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, Any, List
from datetime import datetime
from daily_content_generator import DailyContentGenerator


class PlanDocumentGenerator:
    """学习计划文档生成器"""
    
    def __init__(self):
        self.content_generator = DailyContentGenerator()
        self.output_dir = "word_learning_details"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        
        # 词性映射表
        self.part_of_speech_map = {
            "noun": {"chinese": "名词", "abbreviation": "n."},
            "verb": {"chinese": "动词", "abbreviation": "v."},
            "adjective": {"chinese": "形容词", "abbreviation": "adj."},
            "adverb": {"chinese": "副词", "abbreviation": "adv."},
            "pronoun": {"chinese": "代词", "abbreviation": "pron."},
            "preposition": {"chinese": "介词", "abbreviation": "prep."},
            "conjunction": {"chinese": "连词", "abbreviation": "conj."},
            "interjection": {"chinese": "感叹词", "abbreviation": "interj."},
            "article": {"chinese": "冠词", "abbreviation": "art."},
            "numeral": {"chinese": "数词", "abbreviation": "num."},
            "determiner": {"chinese": "限定词", "abbreviation": "det."}
        }
    
    def _get_part_of_speech_display(self, part_of_speech: str) -> str:
        """获取词性的中文和简写显示"""
        if part_of_speech in self.part_of_speech_map:
            mapping = self.part_of_speech_map[part_of_speech]
            return f"{mapping['chinese']} ({mapping['abbreviation']})"
        else:
            return f"{part_of_speech} ({part_of_speech[0]}. )"
    
    def generate_enhanced_daily_plan(self, day: int) -> str:
        """生成增强版每日学习计划"""
        # 生成每日学习内容
        daily_content = self.content_generator.generate_daily_content(day)
        
        # 创建文档
        doc = Document()
        self._setup_document_styles(doc)
        
        # 添加标题
        doc.add_heading(f"第{day}天学习计划", 0)
        
        # 添加阶段信息
        phase = daily_content["phase"]
        phase_para = doc.add_paragraph(f"学习阶段：{phase['name']}（第{phase['phase']}阶段）")
        phase_para.paragraph_format.line_spacing = 1.2
        
        # 添加学习内容概览
        self._add_learning_overview(doc, daily_content)
        
        # 添加单词学习内容
        self._add_word_learning_content(doc, daily_content["word_content"])
        
        # 添加语法学习内容
        self._add_grammar_learning_content(doc, daily_content["grammar_content"])
        
        # 添加综合练习句子
        self._add_integrated_sentences(doc, daily_content["integrated_sentences"])
        
        # 添加练习题（不包含答案）
        self._add_exercises(doc, daily_content["exercises"], include_answers=False)
        
        # 添加练习题答案（单独一页）
        self._add_exercise_answers(doc, daily_content["exercises"])
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"enhanced_daily_plan_day_{day}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_learning_overview(self, doc: Document, daily_content: Dict[str, Any]):
        """添加学习内容概览"""
        doc.add_heading("学习内容概览", 1)
        
        # 创建概览表格
        overview_table = doc.add_table(rows=4, cols=2)
        overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        overview_table.style = 'Table Grid'
        
        word_content = daily_content["word_content"]
        grammar_content = daily_content["grammar_content"]
        
        overview_data = [
            ("学习时间", "15分钟"),
            ("单词学习", f"{word_content['count']}个单词（{word_content['level']}级别，{word_content['difficulty']}难度）"),
            ("语法学习", f"{grammar_content['topic']}"),
            ("综合练习", f"{len(daily_content['integrated_sentences'])}个综合句子")
        ]
        
        for i, (key, value) in enumerate(overview_data):
            overview_table.cell(i, 0).text = key
            overview_table.cell(i, 1).text = value
            
            # 设置单元格样式
            for j in range(2):
                cell = overview_table.cell(i, j)
                cell.paragraphs[0].runs[0].font.name = '宋体'
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.2
        
        doc.add_paragraph()
    
    def _add_word_learning_content(self, doc: Document, word_content: Dict[str, Any]):
        """添加单词学习内容"""
        doc.add_heading("单词学习内容", 1)
        
        # 添加单词列表
        doc.add_heading("今日单词", 2)
        word_table = doc.add_table(rows=1, cols=5)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        word_table.style = 'Table Grid'
        
        # 设置表头
        header_cells = word_table.rows[0].cells
        headers = ["单词", "发音", "词性", "中文意思", "例句"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.name = '宋体'
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
            header_cells[i].paragraphs[0].paragraph_format.line_spacing = 1.2
        
        # 添加单词数据
        for word_info in word_content["words"]:
            row_cells = word_table.add_row().cells
            row_cells[0].text = word_info.word
            row_cells[1].text = word_info.pronunciation
            # 使用词性的中文和简写显示
            part_of_speech_display = self._get_part_of_speech_display(word_info.part_of_speech)
            row_cells[2].text = part_of_speech_display
            row_cells[3].text = word_info.chinese_meaning
            row_cells[4].text = word_info.example_sentence
            
            # 设置单元格样式
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.name = '宋体'
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.2
        
        doc.add_paragraph()
    
    def _add_grammar_learning_content(self, doc: Document, grammar_content: Dict[str, Any]):
        """添加语法学习内容"""
        doc.add_heading("语法学习内容", 1)
        
        # 添加语法主题
        doc.add_heading(f"语法主题：{grammar_content['topic']}", 2)
        
        # 添加语法规则
        if grammar_content["config"] and "explanation" in grammar_content["config"]:
            doc.add_heading("语法规则", 3)
            rules = grammar_content["config"]["explanation"].get("basic_rules", [])
            for rule in rules:
                rule_para = doc.add_paragraph(f"• {rule}")
                rule_para.paragraph_format.line_spacing = 1.2
        
        # 添加例句
        if grammar_content["config"] and "examples" in grammar_content["config"]:
            doc.add_heading("例句", 3)
            examples = grammar_content["config"]["examples"].get("basic", [])
            for example in examples:
                example_para = doc.add_paragraph(f"• {example}")
                example_para.paragraph_format.line_spacing = 1.2
        
        doc.add_paragraph()
    
    def _add_integrated_sentences(self, doc: Document, integrated_sentences: List[Dict[str, Any]]):
        """添加综合练习句子"""
        doc.add_heading("综合练习句子", 1)
        doc.add_paragraph("以下句子将今日学习的单词和语法结合起来，帮助加强记忆：")
        
        for i, sentence_data in enumerate(integrated_sentences, 1):
            # 添加句子标题（最小间距）
            sentence_heading = doc.add_heading(f"练习句子 {i}", 2)
            sentence_heading.paragraph_format.line_spacing = 1.2
            sentence_heading.paragraph_format.space_after = Pt(2)
            
            # 添加单词信息（最小间距）
            word_para = doc.add_paragraph()
            word_para.add_run("单词：").bold = True
            word_para.add_run(f"{sentence_data['word']} ({sentence_data['word_meaning']}) - {sentence_data['part_of_speech']}")
            word_para.paragraph_format.line_spacing = 1.2
            word_para.paragraph_format.space_after = Pt(0)
            
            # 添加语法主题（最小间距）
            grammar_para = doc.add_paragraph()
            grammar_para.add_run("语法：").bold = True
            grammar_para.add_run(sentence_data['grammar_topic'])
            grammar_para.paragraph_format.line_spacing = 1.2
            grammar_para.paragraph_format.space_after = Pt(0)
            
            # 添加句子（最小间距）
            sentence_para = doc.add_paragraph()
            sentence_para.add_run("英文句子：").bold = True
            sentence_para.add_run(sentence_data['sentence'])
            sentence_para.paragraph_format.line_spacing = 1.2
            sentence_para.paragraph_format.space_after = Pt(0)
            
            # 添加中文翻译（最小间距）
            translation_para = doc.add_paragraph()
            translation_para.add_run("中文翻译：").bold = True
            translation_para.add_run(sentence_data['chinese_translation'])
            translation_para.paragraph_format.line_spacing = 1.2
            translation_para.paragraph_format.space_after = Pt(0)
            
            # 添加语法解释（最小间距）
            explanation_para = doc.add_paragraph()
            explanation_para.add_run("语法解释：").bold = True
            explanation_para.add_run(sentence_data['grammar_explanation'])
            explanation_para.paragraph_format.line_spacing = 1.2
            explanation_para.paragraph_format.space_after = Pt(0)
            
            # 去掉分隔线，最小间隔
    
    def _add_exercises(self, doc: Document, exercises: List[Any], include_answers: bool = True):
        """添加练习题"""
        doc.add_heading("练习题", 1)
        doc.add_paragraph("以下练习题将帮助您巩固今日学习的内容：")
        
        for i, exercise in enumerate(exercises, 1):
            # 添加题目标题（最小间距）
            exercise_heading = doc.add_heading(f"练习 {i} - {self._get_exercise_type_name(exercise.question_type)}", 2)
            exercise_heading.paragraph_format.line_spacing = 1.2
            exercise_heading.paragraph_format.space_after = Pt(1)
            
            # 添加题目（最小间距）
            question_para = doc.add_paragraph()
            question_para.add_run("题目：").bold = True
            question_para.add_run(exercise.question)
            question_para.paragraph_format.line_spacing = 1.2
            question_para.paragraph_format.space_after = Pt(0)
            
            # 如果是选择题，添加选项
            if exercise.question_type == "choice" and exercise.options:
                options_para = doc.add_paragraph()
                options_para.add_run("选项：").bold = True
                options_para.paragraph_format.line_spacing = 1.2
                options_para.paragraph_format.space_after = Pt(0)
                
                for j, option in enumerate(exercise.options, 1):
                    option_para = doc.add_paragraph(f"  {chr(64+j)}. {option}")
                    option_para.paragraph_format.line_spacing = 1.2
                    option_para.paragraph_format.space_after = Pt(0)
            
            # 如果包含答案，添加答案和解释
            if include_answers:
                # 添加答案
                answer_para = doc.add_paragraph()
                answer_para.add_run("答案：").bold = True
                answer_para.add_run(exercise.correct_answer)
                answer_para.paragraph_format.line_spacing = 1.2
                
                # 添加解释
                if exercise.explanation:
                    explanation_para = doc.add_paragraph()
                    explanation_para.add_run("解释：").bold = True
                    explanation_para.add_run(exercise.explanation)
                    explanation_para.paragraph_format.line_spacing = 1.2
            
            # 去掉分隔线，最小间隔
    
    def _get_exercise_type_name(self, question_type: str) -> str:
        """获取题型名称"""
        type_names = {
            "fill_blank": "填空题",
            "translation": "翻译题",
            "choice": "选择题",
            "sentence_completion": "句子完成题"
        }
        return type_names.get(question_type, "练习题")
    
    def _add_exercise_answers(self, doc: Document, exercises: List[Any]):
        """添加练习题答案（单独一页）"""
        # 添加分页符
        doc.add_page_break()
        
        # 添加答案标题
        doc.add_heading("练习题答案", 0)
        doc.add_paragraph("以下是练习题的答案和解释：")
        
        for i, exercise in enumerate(exercises, 1):
            # 添加题目标题（最小间距）
            exercise_heading = doc.add_heading(f"练习 {i} - {self._get_exercise_type_name(exercise.question_type)}", 1)
            exercise_heading.paragraph_format.line_spacing = 1.2
            exercise_heading.paragraph_format.space_after = Pt(2)
            
            # 添加题目（最小间距）
            question_para = doc.add_paragraph()
            question_para.add_run("题目：").bold = True
            question_para.add_run(exercise.question)
            question_para.paragraph_format.line_spacing = 1.2
            question_para.paragraph_format.space_after = Pt(0)
            
            # 如果是选择题，添加选项
            if exercise.question_type == "choice" and exercise.options:
                options_para = doc.add_paragraph()
                options_para.add_run("选项：").bold = True
                options_para.paragraph_format.line_spacing = 1.2
                options_para.paragraph_format.space_after = Pt(0)
                
                for j, option in enumerate(exercise.options, 1):
                    option_para = doc.add_paragraph(f"  {chr(64+j)}. {option}")
                    option_para.paragraph_format.line_spacing = 1.2
                    option_para.paragraph_format.space_after = Pt(0)
            
            # 添加答案（最小间距）
            answer_para = doc.add_paragraph()
            answer_para.add_run("答案：").bold = True
            answer_para.add_run(exercise.correct_answer)
            answer_para.paragraph_format.line_spacing = 1.2
            answer_para.paragraph_format.space_after = Pt(0)
            
            # 添加解释（最小间距）
            if exercise.explanation:
                explanation_para = doc.add_paragraph()
                explanation_para.add_run("解释：").bold = True
                explanation_para.add_run(exercise.explanation)
                explanation_para.paragraph_format.line_spacing = 1.2
                explanation_para.paragraph_format.space_after = Pt(0)
            
            # 去掉分隔线，最小间隔
    
    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.2
        
        # 设置标题样式
        heading1 = doc.styles['Heading 1']
        heading1.font.name = '宋体'
        heading1.font.size = Pt(14)
        heading1.paragraph_format.line_spacing = 1.2
        
        heading2 = doc.styles['Heading 2']
        heading2.font.name = '宋体'
        heading2.font.size = Pt(12)
        heading2.paragraph_format.line_spacing = 1.2
        
        heading3 = doc.styles['Heading 3']
        heading3.font.name = '宋体'
        heading3.font.size = Pt(10.5)
        heading3.paragraph_format.line_spacing = 1.2
