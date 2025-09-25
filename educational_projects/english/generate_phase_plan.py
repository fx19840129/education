#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
按阶段生成英语学习计划
支持生成指定阶段的完整学习计划
"""

import sys
import os
import argparse
from datetime import datetime, timedelta
import json

# 添加模块路径
sys.path.append(os.path.join(os.path.dirname(__file__), 'plan_modules'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'word_learning_modules'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'grammar_modules'))

try:
    from daily_content_generator import DailyContentGenerator
    from plan_document_generator import PlanDocumentGenerator
    from word_database import WordDatabase
    from grammar_config_loader import GrammarConfigLoader
except ImportError as e:
    print(f"❌ 模块导入失败: {e}")
    print("请确保在正确的目录下运行此脚本")
    sys.exit(1)

class PhasePlanGenerator:
    """按阶段学习计划生成器"""
    
    def __init__(self):
        """初始化生成器"""
        print("📚 初始化阶段学习计划生成器...")
        
        self.daily_generator = DailyContentGenerator()
        self.doc_generator = PlanDocumentGenerator()
        self.word_db = WordDatabase()
        self.grammar_loader = GrammarConfigLoader()
        
        # 学习阶段配置
        self.learning_phases = self._setup_learning_phases()
        
        print("✅ 初始化完成")
    
    def _setup_learning_phases(self):
        """设置学习阶段"""
        return {
            "第一阶段 (第1-4周)": {
                "description": "基础语法和高频词汇",
                "weeks": 4,
                "focus": "基础建立",
                "grammar_topics": [
                    "be动词用法-基础用法", "一般现在时-基础用法", "名词单复数-基础规则",
                    "形容词-基础用法", "there_be-基础用法"
                ],
                "word_level": "elementary",
                "daily_words": 8,
                "exercises_per_day": 15
            },
            "第二阶段 (第5-8周)": {
                "description": "进阶语法和词汇扩展",
                "weeks": 4,
                "focus": "技能提升",
                "grammar_topics": [
                    "现在进行时-基础用法", "过去时-基础用法", "形容词比较级-基础规则",
                    "现在完成时-基础用法", "情态动词-基础用法"
                ],
                "word_level": "elementary",
                "daily_words": 10,
                "exercises_per_day": 18
            },
            "第三阶段 (第9-12周)": {
                "description": "复杂语法结构",
                "weeks": 4,
                "focus": "语法深化",
                "grammar_topics": [
                    "被动语态-基础用法", "定语从句-基础用法", "条件句-基础用法",
                    "过去进行时-基础用法", "间接引语-基础用法"
                ],
                "word_level": "middle_school",
                "daily_words": 8,
                "exercises_per_day": 20
            },
            "第四阶段 (第13-16周)": {
                "description": "中级词汇和语法综合",
                "weeks": 4,
                "focus": "综合运用",
                "grammar_topics": [
                    "现在完成时-持续用法", "被动语态-时态变化", "定语从句-关系代词",
                    "条件句-虚拟语气", "非谓语动词-基础用法"
                ],
                "word_level": "middle_school",
                "daily_words": 10,
                "exercises_per_day": 22
            },
            "第五阶段 (第17-20周)": {
                "description": "高级语法和表达",
                "weeks": 4,
                "focus": "表达提升",
                "grammar_topics": [
                    "倒装句-基础用法", "强调句-基础用法", "主谓一致-复杂情况",
                    "虚拟语气-完整体系", "非谓语动词-高级用法"
                ],
                "word_level": "middle_school",
                "daily_words": 12,
                "exercises_per_day": 25
            },
            "第六阶段 (第21-24周)": {
                "description": "综合复习和应用",
                "weeks": 4,
                "focus": "巩固提高",
                "grammar_topics": [
                    "语法综合复习", "错误句型纠正", "高级句型练习",
                    "写作语法应用", "口语语法实践"
                ],
                "word_level": "middle_school",
                "daily_words": 15,
                "exercises_per_day": 30
            }
        }
    
    def list_phases(self):
        """列出所有可用阶段"""
        print("📋 可用学习阶段:")
        print("-" * 50)
        for i, (name, config) in enumerate(self.learning_phases.items(), 1):
            print(f"{i}. {name}")
            print(f"   描述: {config['description']}")
            print(f"   重点: {config['focus']}")
            print(f"   词汇级别: {config['word_level']}")
            print(f"   每日单词: {config['daily_words']}个")
            print(f"   每日练习: {config['exercises_per_day']}题")
            print()
    
    def generate_phase_plan(self, phase_number=1, start_date=None, output_dir="learning_plans"):
        """生成指定阶段的完整学习计划"""
        if start_date is None:
            start_date = datetime.now()
        elif isinstance(start_date, str):
            start_date = datetime.strptime(start_date, "%Y-%m-%d")
        
        # 获取阶段配置
        phase_names = list(self.learning_phases.keys())
        if phase_number < 1 or phase_number > len(phase_names):
            raise ValueError(f"阶段编号必须在1-{len(phase_names)}之间")
        
        phase_name = phase_names[phase_number - 1]
        phase_config = self.learning_phases[phase_name]
        
        print(f"📖 生成{phase_name}学习计划")
        print(f"📅 开始日期: {start_date.strftime('%Y-%m-%d')}")
        print(f"⏱️ 持续时间: {phase_config['weeks']}周")
        print(f"🎯 学习重点: {phase_config['focus']}")
        print(f"📚 词汇级别: {phase_config['word_level']}")
        print(f"📝 每日单词: {phase_config['daily_words']}个")
        print(f"✏️ 每日练习: {phase_config['exercises_per_day']}题")
        
        # 生成阶段计划数据
        phase_data = {
            "metadata": {
                "phase_name": phase_name,
                "phase_number": phase_number,
                "start_date": start_date.strftime("%Y-%m-%d"),
                "end_date": (start_date + timedelta(weeks=phase_config["weeks"]) - timedelta(days=1)).strftime("%Y-%m-%d"),
                "total_days": phase_config["weeks"] * 7,
                "total_weeks": phase_config["weeks"],
                "generation_time": datetime.now().isoformat()
            },
            "phase_info": {
                "name": phase_name,
                "description": phase_config["description"],
                "focus": phase_config["focus"],
                "grammar_topics": phase_config["grammar_topics"],
                "word_level": phase_config["word_level"],
                "daily_words": phase_config["daily_words"],
                "exercises_per_day": phase_config["exercises_per_day"]
            },
            "weeks": [],
            "statistics": {}
        }
        
        # 生成每周详细内容
        total_words = 0
        total_exercises = 0
        current_date = start_date
        
        for week in range(phase_config["weeks"]):
            week_start = current_date + timedelta(weeks=week)
            week_data = self._generate_week_plan(
                week_start, 
                phase_config,
                f"{phase_name} - 第{week+1}周"
            )
            phase_data["weeks"].append(week_data)
            
            total_words += week_data["total_words"]
            total_exercises += week_data["total_exercises"]
        
        # 统计信息
        phase_data["statistics"] = {
            "total_words": total_words,
            "total_exercises": total_exercises,
            "average_words_per_day": round(total_words / (phase_config["weeks"] * 7), 1),
            "average_exercises_per_day": round(total_exercises / (phase_config["weeks"] * 7), 1)
        }
        
        # 保存文件
        files = self._save_phase_plan(phase_data, phase_number, output_dir)
        
        return {
            "phase_data": phase_data,
            "files": files
        }
    
    def _generate_week_plan(self, week_start, phase_config, week_title):
        """生成单周学习计划"""
        week_data = {
            "title": week_title,
            "start_date": week_start.strftime("%Y-%m-%d"),
            "end_date": (week_start + timedelta(days=6)).strftime("%Y-%m-%d"),
            "days": [],
            "total_words": 0,
            "total_exercises": 0
        }
        
        # 为这一周选择语法主题（循环使用）
        grammar_topics = phase_config["grammar_topics"]
        
        for day in range(7):
            day_date = week_start + timedelta(days=day)
            day_name = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][day]
            
            # 周末减少学习量
            if day >= 5:  # 周六、周日
                daily_words = max(3, phase_config["daily_words"] // 2)
                exercises = max(5, phase_config["exercises_per_day"] // 2)
            else:
                daily_words = phase_config["daily_words"]
                exercises = phase_config["exercises_per_day"]
            
            grammar_topic = grammar_topics[day % len(grammar_topics)]
            
            day_data = {
                "date": day_date.strftime("%Y-%m-%d"),
                "day_name": day_name,
                "grammar_topic": grammar_topic,
                "word_count": daily_words,
                "exercise_count": exercises,
                "word_level": phase_config["word_level"]
            }
            
            week_data["days"].append(day_data)
            week_data["total_words"] += daily_words
            week_data["total_exercises"] += exercises
        
        return week_data
    
    def _save_phase_plan(self, phase_data, phase_number, output_dir):
        """保存阶段计划到文件"""
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        phase_name = phase_data["phase_info"]["name"].replace(" ", "_").replace("(", "").replace(")", "")
        
        # 保存JSON格式
        json_file = f"{output_dir}/phase_{phase_number}_{phase_name}_{timestamp}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(phase_data, f, ensure_ascii=False, indent=2)
        
        # 生成Word文档
        word_file = f"{output_dir}/phase_{phase_number}_{phase_name}_{timestamp}.docx"
        self._generate_phase_document(phase_data, word_file)
        
        # 生成文本总结
        txt_file = f"{output_dir}/phase_{phase_number}_{phase_name}_summary_{timestamp}.txt"
        self._generate_phase_summary(phase_data, txt_file)
        
        return {
            "json_file": json_file,
            "word_file": word_file,
            "txt_file": txt_file
        }
    
    def _generate_phase_document(self, phase_data, output_file):
        """生成阶段计划的Word文档"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(f'{phase_data["phase_info"]["name"]}学习计划', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("阶段名称", phase_data["phase_info"]["name"]),
            ("开始日期", phase_data["metadata"]["start_date"]),
            ("结束日期", phase_data["metadata"]["end_date"]),
            ("总天数", f"{phase_data['metadata']['total_days']} 天"),
            ("预期词汇量", f"{phase_data['statistics']['total_words']} 个"),
            ("练习题总数", f"{phase_data['statistics']['total_exercises']} 题")
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # 阶段信息
        doc.add_heading('阶段详情', level=1)
        
        p = doc.add_paragraph()
        p.add_run("学习重点：").bold = True
        p.add_run(f" {phase_data['phase_info']['focus']}")
        
        p = doc.add_paragraph()
        p.add_run("阶段描述：").bold = True
        p.add_run(f" {phase_data['phase_info']['description']}")
        
        p = doc.add_paragraph()
        p.add_run("语法主题：").bold = True
        p.add_run(f" {', '.join(phase_data['phase_info']['grammar_topics'])}")
        
        p = doc.add_paragraph()
        p.add_run("词汇级别：").bold = True
        p.add_run(f" {phase_data['phase_info']['word_level']}")
        
        p = doc.add_paragraph()
        p.add_run("每日单词：").bold = True
        p.add_run(f" {phase_data['phase_info']['daily_words']} 个")
        
        p = doc.add_paragraph()
        p.add_run("每日练习：").bold = True
        p.add_run(f" {phase_data['phase_info']['exercises_per_day']} 题")
        
        # 周计划详情
        doc.add_heading('周计划详情', level=1)
        
        for i, week in enumerate(phase_data["weeks"], 1):
            doc.add_heading(f'第{i}周 ({week["start_date"]} - {week["end_date"]})', level=2)
            
            # 周统计
            p = doc.add_paragraph()
            p.add_run("本周单词：").bold = True
            p.add_run(f" {week['total_words']} 个")
            p.add_run(" | 本周练习：").bold = True
            p.add_run(f" {week['total_exercises']} 题")
            
            # 每日安排
            doc.add_paragraph("每日安排：")
            for day in week["days"]:
                p = doc.add_paragraph(f"  {day['day_name']} ({day['date']}): {day['grammar_topic']} - {day['word_count']}词 + {day['exercise_count']}题", style='List Bullet')
        
        # 学习建议
        doc.add_heading('学习建议', level=1)
        suggestions = [
            f"本阶段重点学习{phase_data['phase_info']['focus']}",
            f"每天坚持学习{phase_data['phase_info']['daily_words']}个单词",
            f"完成{phase_data['phase_info']['exercises_per_day']}道练习题",
            "及时复习前面学过的语法点",
            "注意词汇和语法的结合运用",
            "可以根据个人进度适当调整学习量"
        ]
        
        for suggestion in suggestions:
            p = doc.add_paragraph(suggestion, style='List Bullet')
        
        doc.save(output_file)
    
    def _generate_phase_summary(self, phase_data, output_file):
        """生成阶段计划文本总结"""
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"{phase_data['phase_info']['name']}学习计划总结\n")
            f.write("=" * 50 + "\n\n")
            
            f.write(f"阶段名称: {phase_data['phase_info']['name']}\n")
            f.write(f"开始日期: {phase_data['metadata']['start_date']}\n")
            f.write(f"结束日期: {phase_data['metadata']['end_date']}\n")
            f.write(f"总学习天数: {phase_data['metadata']['total_days']} 天\n")
            f.write(f"预期词汇量: {phase_data['statistics']['total_words']} 个\n")
            f.write(f"练习题总数: {phase_data['statistics']['total_exercises']} 题\n")
            f.write(f"日均单词: {phase_data['statistics']['average_words_per_day']} 个\n")
            f.write(f"日均练习: {phase_data['statistics']['average_exercises_per_day']} 题\n\n")
            
            f.write("阶段详情:\n")
            f.write("-" * 30 + "\n")
            f.write(f"学习重点: {phase_data['phase_info']['focus']}\n")
            f.write(f"阶段描述: {phase_data['phase_info']['description']}\n")
            f.write(f"词汇级别: {phase_data['phase_info']['word_level']}\n")
            f.write(f"每日单词: {phase_data['phase_info']['daily_words']} 个\n")
            f.write(f"每日练习: {phase_data['phase_info']['exercises_per_day']} 题\n\n")
            
            f.write("语法主题:\n")
            for topic in phase_data['phase_info']['grammar_topics']:
                f.write(f"  - {topic}\n")
            
            f.write("\n周计划概览:\n")
            f.write("-" * 30 + "\n")
            
            for i, week in enumerate(phase_data["weeks"], 1):
                f.write(f"\n第{i}周 ({week['start_date']} - {week['end_date']}):\n")
                f.write(f"  单词: {week['total_words']} 个\n")
                f.write(f"  练习: {week['total_exercises']} 题\n")
                f.write("  每日安排:\n")
                for day in week["days"]:
                    f.write(f"    {day['day_name']}: {day['grammar_topic']} ({day['word_count']}词 + {day['exercise_count']}题)\n")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description="按阶段生成英语学习计划")
    parser.add_argument("--phase", "-p", type=int, help="生成指定阶段的完整学习计划 (1-6)")
    parser.add_argument("--start-date", "-s", help="开始日期 (YYYY-MM-DD)")
    parser.add_argument("--output-dir", "-o", default="learning_plans", 
                       help="输出目录")
    parser.add_argument("--list-phases", action="store_true", help="列出所有可用阶段")
    
    args = parser.parse_args()
    
    print("📚 === 按阶段英语学习计划生成器 ===")
    print("🎯 支持生成指定阶段的完整学习计划")
    print()
    
    try:
        generator = PhasePlanGenerator()
        
        if args.list_phases:
            # 列出所有阶段
            generator.list_phases()
            
        elif args.phase:
            # 生成指定阶段计划
            result = generator.generate_phase_plan(
                phase_number=args.phase,
                start_date=args.start_date,
                output_dir=args.output_dir
            )
            
            phase_data = result["phase_data"]
            files = result["files"]
            
            print(f"\n🎉 {phase_data['phase_info']['name']}学习计划生成完成！")
            print(f"📄 详细计划: {files['word_file']}")
            print(f"📊 数据文件: {files['json_file']}")
            print(f"📋 总结文件: {files['txt_file']}")
            
            print(f"\n📈 阶段统计:")
            stats = phase_data["statistics"]
            print(f"   总词汇量: {stats['total_words']} 个")
            print(f"   总练习量: {stats['total_exercises']} 题")
            print(f"   日均单词: {stats['average_words_per_day']} 个")
            print(f"   日均练习: {stats['average_exercises_per_day']} 题")
            print(f"   学习天数: {phase_data['metadata']['total_days']} 天")
            print(f"   学习周数: {phase_data['metadata']['total_weeks']} 周")
            
        else:
            # 默认列出所有阶段
            print("💡 请指定要生成的阶段编号，或使用 --list-phases 查看所有可用阶段")
            print("\n使用方法:")
            print("  python generate_phase_plan.py --list-phases  # 列出所有阶段")
            print("  python generate_phase_plan.py --phase 1      # 生成第一阶段计划")
            print("  python generate_phase_plan.py --phase 2 -s 2024-01-01  # 指定开始日期")
    
    except Exception as e:
        print(f"❌ 生成失败: {e}")
        print("💡 请检查模块是否正确安装")


if __name__ == "__main__":
    main()
