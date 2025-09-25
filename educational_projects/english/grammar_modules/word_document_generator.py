#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档生成器
生成包含语法讲解、例句、练习题的Word文档
"""

import os
import datetime
from typing import Dict, List, Any
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


class WordDocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self, output_dir: str = "word_grammar_details"):
        """
        初始化Word文档生成器
        
        Args:
            output_dir: 输出目录
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_grammar_word_document(self, grammar_config: Dict[str, Any], 
                                     exercises: List[Dict[str, Any]]) -> tuple:
        """
        生成语法Word文档
        
        Args:
            grammar_config: 语法配置
            exercises: 练习题列表
            
        Returns:
            (主文档路径, 答案文档路径)
        """
        grammar_name = grammar_config.get("grammar_name", "")
        level = grammar_config.get("level", "")
        category = grammar_config.get("category", "")
        
        # 创建主文档
        main_doc = self._create_main_document(grammar_config, exercises)
        main_file = self.output_dir / f"{grammar_name}_语法详解_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        main_doc.save(str(main_file))
        
        # 创建答案文档
        answer_doc = self._create_answer_document(grammar_config, exercises)
        answer_file = self.output_dir / f"{grammar_name}_练习题答案_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        answer_doc.save(str(answer_file))
        
        return str(main_file), str(answer_file)
    
    def _create_main_document(self, grammar_config: Dict[str, Any], 
                            exercises: List[Dict[str, Any]]) -> Document:
        """创建主文档"""
        doc = Document()
        
        # 设置文档样式
        self._setup_document_styles(doc)
        
        # 添加标题
        title = doc.add_heading(f"{grammar_config.get('grammar_name', '')} - 详细语法讲解", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        self._add_basic_info(doc, grammar_config)
        
        # 添加语法讲解
        self._add_grammar_explanation(doc, grammar_config)
        
        # 添加例句
        self._add_examples(doc, grammar_config)
        
        # 添加练习题
        self._add_exercises(doc, exercises)
        
        # 添加学习目标
        self._add_learning_objectives(doc, grammar_config)
        
        # 添加评估标准
        self._add_assessment_criteria(doc, grammar_config)
        
        return doc
    
    def _create_answer_document(self, grammar_config: Dict[str, Any], 
                              exercises: List[Dict[str, Any]]) -> Document:
        """创建答案文档"""
        doc = Document()
        
        # 设置文档样式
        self._setup_document_styles(doc)
        
        # 添加标题
        title = doc.add_heading(f"{grammar_config.get('grammar_name', '')} - 练习题答案", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        self._add_basic_info(doc, grammar_config)
        
        # 添加答案详解
        self._add_answer_details(doc, exercises)
        
        return doc
    
    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置页边距（节省纸张）
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)      # 上边距
            section.bottom_margin = Inches(0.5)   # 下边距
            section.left_margin = Inches(0.6)     # 左边距
            section.right_margin = Inches(0.6)    # 右边距
        
        # 设置默认字体（减小字体节省空间）
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)  # 从12减小到10.5
        
        # 创建标题样式
        heading_style = doc.styles['Heading 1']
        heading_font = heading_style.font
        heading_font.name = '黑体'
        heading_font.size = Pt(14)  # 从16减小到14
        
        # 创建二级标题样式
        heading2_style = doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = '黑体'
        heading2_font.size = Pt(12)  # 从14减小到12
        
        # 创建三级标题样式
        heading3_style = doc.styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = '黑体'
        heading3_font.size = Pt(10.5)  # 从12减小到10.5
    
    def _add_basic_info(self, doc: Document, grammar_config: Dict[str, Any]):
        """添加基本信息"""
        # 添加分隔线
        doc.add_paragraph("=" * 50)
        
        # 适用年级
        level_para = doc.add_paragraph()
        level_para.add_run("适用年级：").bold = True
        level_para.add_run(grammar_config.get('level', ''))
        level_para.paragraph_format.line_spacing = 1.2
        
        # 难度级别
        difficulty_para = doc.add_paragraph()
        difficulty_para.add_run("难度级别：").bold = True
        difficulty_para.add_run(grammar_config.get('difficulty', '').upper())
        difficulty_para.paragraph_format.line_spacing = 1.2
        
        # 语法分类
        category_para = doc.add_paragraph()
        category_para.add_run("语法分类：").bold = True
        category_para.add_run(grammar_config.get('category', ''))
        category_para.paragraph_format.line_spacing = 1.2
        
        # 语法描述
        desc_para = doc.add_paragraph()
        desc_para.add_run("语法描述：").bold = True
        desc_para.add_run(grammar_config.get('description', ''))
        desc_para.paragraph_format.line_spacing = 1.2
        
        # 添加分隔线
        doc.add_paragraph("=" * 50)
    
    def _add_grammar_explanation(self, doc: Document, grammar_config: Dict[str, Any]):
        """添加语法讲解"""
        doc.add_heading("语法讲解", level=1)
        
        explanation = grammar_config.get('explanation', {})
        
        # 基本规则
        if 'basic_rules' in explanation:
            doc.add_heading("基本规则", level=2)
            for i, rule in enumerate(explanation['basic_rules'], 1):
                para = doc.add_paragraph(f"{i}. {rule}")
                para.style = 'List Number'
                para.paragraph_format.line_spacing = 1.2
        
        # 其他规则
        for key, value in explanation.items():
            if key not in ['basic_rules', 'common_errors', 'usage_tips'] and isinstance(value, str):
                doc.add_heading(key.replace('_', ' ').title(), level=2)
                para = doc.add_paragraph(value)
                para.paragraph_format.line_spacing = 1.2
        
        # 常见错误
        if 'common_errors' in explanation:
            doc.add_heading("常见错误", level=2)
            for i, error in enumerate(explanation['common_errors'], 1):
                para = doc.add_paragraph(f"{i}. {error}")
                para.style = 'List Number'
                para.paragraph_format.line_spacing = 1.2
        
        # 使用技巧
        if 'usage_tips' in explanation:
            doc.add_heading("使用技巧", level=2)
            for i, tip in enumerate(explanation['usage_tips'], 1):
                para = doc.add_paragraph(f"{i}. {tip}")
                para.style = 'List Number'
                para.paragraph_format.line_spacing = 1.2
    
    def _add_examples(self, doc: Document, grammar_config: Dict[str, Any]):
        """添加例句"""
        doc.add_heading("例句", level=1)
        
        examples = grammar_config.get('examples', {})
        for category, example_list in examples.items():
            if isinstance(example_list, list) and example_list:
                doc.add_heading(category.replace('_', ' ').title(), level=2)
                for i, example in enumerate(example_list, 1):
                    para = doc.add_paragraph(f"{i}. {example}")
                    para.style = 'List Number'
                    para.paragraph_format.line_spacing = 1.2
    
    def _add_exercises(self, doc: Document, exercises: List[Dict[str, Any]]):
        """添加练习题"""
        doc.add_heading(f"练习题（共{len(exercises)}道）", level=1)
        
        for i, exercise in enumerate(exercises, 1):
            # 添加练习题标题（减小标题间距）
            exercise_title = doc.add_heading(f"第{i}题 [{exercise['type']}]", level=2)
            exercise_title.paragraph_format.space_after = Pt(6)  # 减小标题后间距
            
            # 添加题目
            question_para = doc.add_paragraph()
            question_para.add_run("题目：").bold = True
            question_para.add_run(exercise['question'])
            # 设置行间距为1.2倍
            question_para.paragraph_format.line_spacing = 1.2
            question_para.paragraph_format.space_after = Pt(3)  # 减小段落后间距
            
            # 添加选项（如果有）
            if exercise.get('options'):
                options_para = doc.add_paragraph()
                options_para.add_run("选项：").bold = True
                options_para.paragraph_format.line_spacing = 1.2
                options_para.paragraph_format.space_after = Pt(3)
                for option in exercise['options']:
                    option_para = doc.add_paragraph(option, style='List Bullet')
                    option_para.paragraph_format.line_spacing = 1.2
                    option_para.paragraph_format.space_after = Pt(1)
            
            # 减少空行
            if i < len(exercises):  # 不是最后一题才添加空行
                doc.add_paragraph()
    
    def _add_answer_details(self, doc: Document, exercises: List[Dict[str, Any]]):
        """添加答案详解"""
        doc.add_heading("答案详解", level=1)
        
        for i, exercise in enumerate(exercises, 1):
            # 添加答案标题（减小标题间距）
            answer_title = doc.add_heading(f"第{i}题答案", level=2)
            answer_title.paragraph_format.space_after = Pt(6)  # 减小标题后间距
            
            # 添加题型
            type_para = doc.add_paragraph()
            type_para.add_run("题型：").bold = True
            type_para.add_run(exercise['type'])
            type_para.paragraph_format.line_spacing = 1.2
            type_para.paragraph_format.space_after = Pt(2)  # 减小段落后间距
            
            # 添加题目
            question_para = doc.add_paragraph()
            question_para.add_run("题目：").bold = True
            question_para.add_run(exercise['question'])
            question_para.paragraph_format.line_spacing = 1.2
            question_para.paragraph_format.space_after = Pt(2)
            
            # 添加答案
            answer_para = doc.add_paragraph()
            answer_para.add_run("答案：").bold = True
            answer_para.add_run(exercise['answer'])
            answer_para.paragraph_format.line_spacing = 1.2
            answer_para.paragraph_format.space_after = Pt(2)
            
            # 添加解析
            explanation_para = doc.add_paragraph()
            explanation_para.add_run("解析：").bold = True
            explanation_para.add_run(exercise['explanation'])
            explanation_para.paragraph_format.line_spacing = 1.2
            explanation_para.paragraph_format.space_after = Pt(2)
            
            # 添加难度
            difficulty_para = doc.add_paragraph()
            difficulty_para.add_run("难度：").bold = True
            difficulty_para.add_run(exercise.get('difficulty', 'medium'))
            difficulty_para.paragraph_format.line_spacing = 1.2
            difficulty_para.paragraph_format.space_after = Pt(2)
            
            # 减少分隔线长度
            if i < len(exercises):  # 不是最后一题才添加分隔线
                doc.add_paragraph("-" * 20)
    
    def _add_learning_objectives(self, doc: Document, grammar_config: Dict[str, Any]):
        """添加学习目标"""
        if 'learning_objectives' in grammar_config:
            doc.add_heading("学习目标", level=1)
            for i, objective in enumerate(grammar_config['learning_objectives'], 1):
                para = doc.add_paragraph(f"{i}. {objective}")
                para.style = 'List Number'
                para.paragraph_format.line_spacing = 1.2
                para.paragraph_format.space_after = Pt(1)
    
    def _add_assessment_criteria(self, doc: Document, grammar_config: Dict[str, Any]):
        """添加评估标准"""
        if 'assessment_criteria' in grammar_config:
            doc.add_heading("评估标准", level=1)
            for level, criteria in grammar_config['assessment_criteria'].items():
                para = doc.add_paragraph()
                para.add_run(f"{level}：").bold = True
                para.add_run(criteria)
                para.paragraph_format.line_spacing = 1.2
                para.paragraph_format.space_after = Pt(1)
    
    def generate_all_grammar_word_documents(self, configs: Dict[str, Dict[str, Any]], 
                                          exercises_dict: Dict[str, List[Dict[str, Any]]]) -> bool:
        """
        生成所有语法点的Word文档
        
        Args:
            configs: 所有语法配置
            exercises_dict: 所有练习题
            
        Returns:
            是否生成成功
        """
        success_count = 0
        
        for grammar_name, config in configs.items():
            exercises = exercises_dict.get(grammar_name, [])
            try:
                main_file, answer_file = self.generate_grammar_word_document(config, exercises)
                print(f"✓ {grammar_name} Word文档已生成")
                print(f"  📄 主文档: {main_file}")
                print(f"  📋 答案文档: {answer_file}")
                success_count += 1
            except Exception as e:
                print(f"❌ {grammar_name} Word文档生成失败: {e}")
        
        print(f"\n✓ 共生成 {success_count}/{len(configs)} 个语法Word文档")
        return success_count > 0


if __name__ == "__main__":
    # 测试Word文档生成器
    from grammar_config_loader import GrammarConfigLoader
    from exercise_generator import ImprovedExerciseGenerator
    
    # 初始化组件
    config_loader = GrammarConfigLoader()
    exercise_generator = ImprovedExerciseGenerator()
    word_generator = WordDocumentGenerator()
    
    # 加载配置
    config = config_loader.load_grammar_config("be动词用法", "elementary")
    if config:
        # 生成练习题
        exercises = exercise_generator.generate_exercises(config, 5, "easy")
        
        # 生成Word文档
        main_file, answer_file = word_generator.generate_grammar_word_document(config, exercises)
        print(f"Word文档已生成：{main_file}, {answer_file}")
    else:
        print("未找到语法配置")
